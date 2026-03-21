import streamlit as st
import pandas as pd
import streamlit.components.v1 as components

st.set_page_config(page_title="AWS D1.1 WPS Reference", page_icon="📋", layout="wide")

st.title("📋 AWS D1.1/D1.1M:2025 — Prequalified WPS Requirements")
st.markdown("Interactive reference for Clause 5 — select a process and filter by position.")
st.divider()

# ════════════════════════════════════════════════════════
# ALL PREVIOUS DATA (keep exactly as before)
# ════════════════════════════════════════════════════════

smaw_data = [
    {"Variable": "Electrode Diameter", "Position": "Flat", "Weld Type": "Fillet", "Requirement": "5/16 in [8.0 mm] max.", "Exceptions": "Except root pass"},
    {"Variable": "Electrode Diameter", "Position": "Flat", "Weld Type": "Groove", "Requirement": "1/4 in [6.4 mm] max.", "Exceptions": "Except root pass"},
    {"Variable": "Electrode Diameter", "Position": "Flat", "Weld Type": "Root Pass", "Requirement": "3/16 in [4.8 mm] max.", "Exceptions": ""},
    {"Variable": "Electrode Diameter", "Position": "Horizontal", "Weld Type": "Fillet", "Requirement": "1/4 in [6.4 mm] max.", "Exceptions": ""},
    {"Variable": "Electrode Diameter", "Position": "Horizontal", "Weld Type": "Groove", "Requirement": "3/16 in [4.8 mm] max.", "Exceptions": ""},
    {"Variable": "Electrode Diameter", "Position": "Vertical", "Weld Type": "All", "Requirement": "3/16 in [4.8 mm] max.", "Exceptions": "5/32 in [4.0 mm] max for EXX14 and low-hydrogen electrodes"},
    {"Variable": "Electrode Diameter", "Position": "Overhead", "Weld Type": "All", "Requirement": "3/16 in [4.8 mm] max.", "Exceptions": "5/32 in [4.0 mm] max for EXX14 and low-hydrogen electrodes"},
    {"Variable": "Current", "Position": "All", "Weld Type": "All", "Requirement": "Within range recommended by filler metal manufacturer", "Exceptions": ""},
    {"Variable": "Root Bead Thickness", "Position": "Flat", "Weld Type": "Groove/Fillet", "Requirement": "3/8 in [10 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Root Bead Thickness", "Position": "Horizontal", "Weld Type": "Groove/Fillet", "Requirement": "5/16 in [8 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Root Bead Thickness", "Position": "Vertical", "Weld Type": "Groove/Fillet", "Requirement": "1/2 in [12 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Root Bead Thickness", "Position": "Overhead", "Weld Type": "Groove/Fillet", "Requirement": "5/16 in [8 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Fill & Cap Bead Thickness", "Position": "All", "Weld Type": "Groove/Fillet", "Requirement": "3/16 in [5 mm] max.", "Exceptions": ""},
    {"Variable": "Single Pass Fillet Size", "Position": "Flat", "Weld Type": "Fillet", "Requirement": "3/8 in [10 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Single Pass Fillet Size", "Position": "Horizontal", "Weld Type": "Fillet", "Requirement": "5/16 in [8 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Single Pass Fillet Size", "Position": "Vertical", "Weld Type": "Fillet", "Requirement": "1/2 in [12 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Single Pass Fillet Size", "Position": "Overhead", "Weld Type": "Fillet", "Requirement": "5/16 in [8 mm] max.", "Exceptions": "See 5.6.2"},
]

saw_data = [
    {"Variable": "Electrode Diameter", "Position": "Flat", "Weld Type": "All", "Single Electrode": "1/4 in [6.4 mm] max.", "Parallel Electrodes": "—", "Multiple Electrodes": "—", "Exceptions": ""},
    {"Variable": "Electrode Diameter", "Position": "Horizontal", "Weld Type": "Fillet", "Single Electrode": "1/4 in [6.4 mm] max.", "Parallel Electrodes": "—", "Multiple Electrodes": "—", "Exceptions": ""},
    {"Variable": "Current — Fillet", "Position": "Flat & Horizontal", "Weld Type": "Fillet", "Single Electrode": "1000 A max.", "Parallel Electrodes": "1200 A max.", "Multiple Electrodes": "Unlimited", "Exceptions": ""},
    {"Variable": "Current — Groove root (with opening)", "Position": "Flat & Horizontal", "Weld Type": "Groove", "Single Electrode": "600 A max.", "Parallel Electrodes": "700 A max.", "Multiple Electrodes": "—", "Exceptions": ""},
    {"Variable": "Current — Groove root (no opening)", "Position": "Flat & Horizontal", "Weld Type": "Groove", "Single Electrode": "900 A max.", "Parallel Electrodes": "—", "Multiple Electrodes": "—", "Exceptions": ""},
    {"Variable": "Current — Groove fill beads", "Position": "Flat & Horizontal", "Weld Type": "Groove", "Single Electrode": "1200 A max.", "Parallel Electrodes": "—", "Multiple Electrodes": "—", "Exceptions": ""},
    {"Variable": "Current — Groove cap bead", "Position": "Flat & Horizontal", "Weld Type": "Groove", "Single Electrode": "Unlimited", "Parallel Electrodes": "—", "Multiple Electrodes": "Unlimited", "Exceptions": ""},
    {"Variable": "Root Bead Thickness", "Position": "Flat & Horizontal", "Weld Type": "All", "Single Electrode": "Unlimited", "Parallel Electrodes": "Unlimited", "Multiple Electrodes": "Unlimited", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Fill & Cap Bead Thickness", "Position": "Flat & Horizontal", "Weld Type": "All", "Single Electrode": "1/4 in [6 mm] max.", "Parallel Electrodes": "1/4 in [6 mm] max.", "Multiple Electrodes": "Unlimited", "Exceptions": ""},
    {"Variable": "Single Pass Fillet Size", "Position": "Flat & Horizontal", "Weld Type": "Fillet", "Single Electrode": "5/16 in [8 mm] max.", "Parallel Electrodes": "5/16 in [8 mm] max.", "Multiple Electrodes": "1/2 in [12 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Layer Width — Root", "Position": "Flat & Horizontal", "Weld Type": "Root Bead", "Single Electrode": "Split root layer", "Parallel Electrodes": "Split root layer", "Multiple Electrodes": "Split root layer", "Exceptions": "For groove opening > 1/2 in [12 mm]"},
    {"Variable": "Layer Width — Fill & cap", "Position": "Flat & Horizontal", "Weld Type": "All", "Single Electrode": "Split if w > 5/8 in, max 1 in", "Parallel Electrodes": "Split if w > 5/8 in, max 1-1/8 in", "Multiple Electrodes": "Split if w > 1 in, max 1-1/4 in", "Exceptions": ""},
]

gmaw_data = [
    {"Variable": "Electrode Diameter", "Position": "All", "Weld Type": "All", "Requirement": "0.0625 in [1.6 mm] max.", "Exceptions": ""},
    {"Variable": "Current (min. amperage)", "Position": "All", "Weld Type": "All", "Requirement": "0.030→190A | 0.035→210A | 0.040→230A | 0.045→260A | 0.0625→300A min.", "Exceptions": "Manufacturer recs may be used if not short circuit"},
    {"Variable": "Pulsed Spray (min.)", "Position": "All", "Weld Type": "All", "Requirement": "0.035→100A | 0.045→110A | 0.052→120A | 0.0625→140A min.", "Exceptions": "Manufacturer recommendations may be used"},
    {"Variable": "Root Bead Thickness", "Position": "Flat", "Weld Type": "Fillet/Groove", "Requirement": "3/8 in [10 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Root Bead Thickness", "Position": "Horizontal", "Weld Type": "Fillet/Groove", "Requirement": "5/16 in [8 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Root Bead Thickness", "Position": "Vertical", "Weld Type": "Fillet/Groove", "Requirement": "1/2 in [12 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Root Bead Thickness", "Position": "Overhead", "Weld Type": "Fillet/Groove", "Requirement": "5/16 in [8 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Fill & Cap Bead Thickness", "Position": "All", "Weld Type": "Fillet/Groove", "Requirement": "1/4 in [6 mm] max.", "Exceptions": ""},
    {"Variable": "Single Pass Fillet Size", "Position": "Flat", "Weld Type": "Fillet", "Requirement": "1/2 in [12 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Single Pass Fillet Size", "Position": "Horizontal", "Weld Type": "Fillet", "Requirement": "3/8 in [10 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Single Pass Fillet Size", "Position": "Vertical", "Weld Type": "Fillet", "Requirement": "1/2 in [12 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Single Pass Fillet Size", "Position": "Overhead", "Weld Type": "Fillet", "Requirement": "5/16 in [8 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Single Pass Layer Width", "Position": "Flat/Horizontal/Overhead", "Weld Type": "Groove/Fillet", "Requirement": "Split layers if w > 5/8 in [16 mm]", "Exceptions": ""},
    {"Variable": "Single Pass Layer Width", "Position": "Vertical", "Weld Type": "Groove/Fillet", "Requirement": "Split layers if w > 1 in [25 mm]", "Exceptions": ""},
]

fcaw_data = [
    {"Variable": "Electrode Dia — FCAW", "Position": "Flat/Horizontal", "T Condition": "T ≤ 3/8 in", "Requirement": "0.030 in [0.8 mm] min.", "Exceptions": ""},
    {"Variable": "Electrode Dia — GMAW Cored", "Position": "Flat/Horizontal", "T Condition": "T ≤ 3/8 in", "Requirement": "No minimum", "Exceptions": ""},
    {"Variable": "Electrode Dia — FCAW & GMAW Cored", "Position": "Flat/Horizontal", "T Condition": "T > 3/8 in", "Requirement": "0.045 in [1 mm] min.", "Exceptions": ""},
    {"Variable": "Electrode Dia max — FCAW", "Position": "Flat/Horizontal", "T Condition": "T > 3/8 in", "Requirement": "1/8 in [3.2 mm] max.", "Exceptions": ""},
    {"Variable": "Electrode Dia — FCAW", "Position": "Vertical", "T Condition": "T ≤ 3/8 in", "Requirement": "0.030 in [0.8 mm] min.", "Exceptions": ""},
    {"Variable": "Electrode Dia max — FCAW", "Position": "Vertical", "T Condition": "T > 3/8 in", "Requirement": "3/32 in [2.4 mm] max.", "Exceptions": ""},
    {"Variable": "Electrode Dia max — FCAW", "Position": "Overhead", "T Condition": "T > 3/8 in", "Requirement": "5/64 in [2.2 mm] max.", "Exceptions": ""},
    {"Variable": "Current", "Position": "All", "T Condition": "All", "Requirement": "Within range recommended by filler metal manufacturer", "Exceptions": ""},
    {"Variable": "Root Bead Thickness", "Position": "Flat", "T Condition": "—", "Requirement": "3/8 in [10 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Root Bead Thickness", "Position": "Horizontal", "T Condition": "—", "Requirement": "5/16 in [8 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Root Bead Thickness", "Position": "Vertical", "T Condition": "—", "Requirement": "1/2 in [12 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Root Bead Thickness", "Position": "Overhead", "T Condition": "—", "Requirement": "5/16 in [8 mm] max.", "Exceptions": "See 5.8.2.1"},
    {"Variable": "Fill & Cap Bead Thickness", "Position": "All", "T Condition": "—", "Requirement": "1/4 in [6 mm] max.", "Exceptions": ""},
    {"Variable": "Single Pass Fillet Size", "Position": "Flat", "T Condition": "—", "Requirement": "1/2 in [12 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Single Pass Fillet Size", "Position": "Horizontal", "T Condition": "—", "Requirement": "3/8 in [10 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Single Pass Fillet Size", "Position": "Vertical", "T Condition": "—", "Requirement": "1/2 in [12 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Single Pass Fillet Size", "Position": "Overhead", "T Condition": "—", "Requirement": "5/16 in [8 mm] max.", "Exceptions": "See 5.6.2"},
    {"Variable": "Layer Width — Groove root opening ≥ 1/2 in", "Position": "All", "T Condition": "—", "Requirement": "Split root layer, max bead 1 in [25 mm]", "Exceptions": ""},
    {"Variable": "Layer Width — Fill & cap", "Position": "Flat/Horizontal/Overhead", "T Condition": "—", "Requirement": "Split if w > 5/8 in [16 mm]", "Exceptions": ""},
    {"Variable": "Layer Width — Fill & cap", "Position": "Vertical", "T Condition": "—", "Requirement": "Split if w > 1 in [25 mm]", "Exceptions": ""},
]

preheat_data = [
    {"Category": "A", "Process": "SMAW (non low-hydrogen)", "Thickness": "1/8 to 3/4 in incl.", "Min Preheat °F": "32°F", "Min Preheat °C": "0°C", "Notes": "If base metal < 32°F, preheat to 70°F min"},
    {"Category": "A", "Process": "SMAW (non low-hydrogen)", "Thickness": "Over 3/4 to 1-1/2 in", "Min Preheat °F": "150°F", "Min Preheat °C": "65°C", "Notes": ""},
    {"Category": "A", "Process": "SMAW (non low-hydrogen)", "Thickness": "Over 1-1/2 to 2-1/2 in", "Min Preheat °F": "225°F", "Min Preheat °C": "110°C", "Notes": ""},
    {"Category": "A", "Process": "SMAW (non low-hydrogen)", "Thickness": "Over 2-1/2 in", "Min Preheat °F": "300°F", "Min Preheat °C": "150°C", "Notes": ""},
    {"Category": "B", "Process": "SMAW (low-hydrogen), SAW, GMAW, FCAW", "Thickness": "1/8 to 3/4 in incl.", "Min Preheat °F": "32°F", "Min Preheat °C": "0°C", "Notes": ""},
    {"Category": "B", "Process": "SMAW (low-hydrogen), SAW, GMAW, FCAW", "Thickness": "Over 3/4 to 1-1/2 in", "Min Preheat °F": "50°F", "Min Preheat °C": "10°C", "Notes": ""},
    {"Category": "B", "Process": "SMAW (low-hydrogen), SAW, GMAW, FCAW", "Thickness": "Over 1-1/2 to 2-1/2 in", "Min Preheat °F": "150°F", "Min Preheat °C": "65°C", "Notes": ""},
    {"Category": "B", "Process": "SMAW (low-hydrogen), SAW, GMAW, FCAW", "Thickness": "Over 2-1/2 in", "Min Preheat °F": "225°F", "Min Preheat °C": "110°C", "Notes": ""},
    {"Category": "C", "Process": "SMAW (low-hydrogen), SAW, GMAW, FCAW", "Thickness": "1/8 to 3/4 in incl.", "Min Preheat °F": "50°F", "Min Preheat °C": "10°C", "Notes": ""},
    {"Category": "C", "Process": "SMAW (low-hydrogen), SAW, GMAW, FCAW", "Thickness": "Over 3/4 to 1-1/2 in", "Min Preheat °F": "150°F", "Min Preheat °C": "65°C", "Notes": ""},
    {"Category": "C", "Process": "SMAW (low-hydrogen), SAW, GMAW, FCAW", "Thickness": "Over 1-1/2 to 2-1/2 in", "Min Preheat °F": "225°F", "Min Preheat °C": "110°C", "Notes": ""},
    {"Category": "C", "Process": "SMAW (low-hydrogen), SAW, GMAW, FCAW", "Thickness": "Over 2-1/2 in", "Min Preheat °F": "300°F", "Min Preheat °C": "150°C", "Notes": ""},
]

pjp_size_data = [
    {"Base Metal Thickness (in)": "1/8 to 3/16 incl.", "Base Metal Thickness (mm)": "3 to 5 incl.", "Min Weld Size (in)": "1/16", "Min Weld Size (mm)": "2"},
    {"Base Metal Thickness (in)": "Over 3/16 to 1/4 incl.", "Base Metal Thickness (mm)": "Over 5 to 6 incl.", "Min Weld Size (in)": "1/8", "Min Weld Size (mm)": "3"},
    {"Base Metal Thickness (in)": "Over 1/4 to 1/2 incl.", "Base Metal Thickness (mm)": "Over 6 to 12 incl.", "Min Weld Size (in)": "3/16", "Min Weld Size (mm)": "5"},
    {"Base Metal Thickness (in)": "Over 1/2 to 3/4 incl.", "Base Metal Thickness (mm)": "Over 12 to 20 incl.", "Min Weld Size (in)": "1/4", "Min Weld Size (mm)": "6"},
    {"Base Metal Thickness (in)": "Over 3/4 to 1-1/2 incl.", "Base Metal Thickness (mm)": "Over 20 to 38 incl.", "Min Weld Size (in)": "5/16", "Min Weld Size (mm)": "8"},
    {"Base Metal Thickness (in)": "Over 1-1/2 to 2-1/4 incl.", "Base Metal Thickness (mm)": "Over 38 to 57 incl.", "Min Weld Size (in)": "3/8", "Min Weld Size (mm)": "10"},
    {"Base Metal Thickness (in)": "Over 2-1/4 to 6 incl.", "Base Metal Thickness (mm)": "Over 57 to 150 incl.", "Min Weld Size (in)": "1/2", "Min Weld Size (mm)": "12"},
    {"Base Metal Thickness (in)": "Over 6", "Base Metal Thickness (mm)": "Over 150", "Min Weld Size (in)": "5/8", "Min Weld Size (mm)": "16"},
]

# ════════════════════════════════════════════════════════
# NEW: TABLE 5.6 — BASE METALS
# ════════════════════════════════════════════════════════
base_metals = [
    # GROUP I
    {"Group": "I", "Specification": "ASTM A36", "Grade/Condition": "≤3/4 in [20 mm]", "Min Yield (ksi)": 36, "Min Yield (MPa)": 250, "Tensile Range (ksi)": "58–80", "Tensile Range (MPa)": "400–550"},
    {"Group": "I", "Specification": "ASTM A53", "Grade/Condition": "Grade B", "Min Yield (ksi)": 35, "Min Yield (MPa)": 240, "Tensile Range (ksi)": "60 min.", "Tensile Range (MPa)": "415 min."},
    {"Group": "I", "Specification": "ASTM A106", "Grade/Condition": "Grade B", "Min Yield (ksi)": 35, "Min Yield (MPa)": 240, "Tensile Range (ksi)": "60 min.", "Tensile Range (MPa)": "415 min."},
    {"Group": "I", "Specification": "ASTM A131", "Grade/Condition": "Grades A, B, D, E", "Min Yield (ksi)": 34, "Min Yield (MPa)": 235, "Tensile Range (ksi)": "58–75", "Tensile Range (MPa)": "400–520"},
    {"Group": "I", "Specification": "ASTM A500", "Grade/Condition": "Grade B", "Min Yield (ksi)": 46, "Min Yield (MPa)": 315, "Tensile Range (ksi)": "58 min.", "Tensile Range (MPa)": "400 min."},
    {"Group": "I", "Specification": "ASTM A500", "Grade/Condition": "Grade C", "Min Yield (ksi)": 50, "Min Yield (MPa)": 345, "Tensile Range (ksi)": "62 min.", "Tensile Range (MPa)": "425 min."},
    {"Group": "I", "Specification": "ASTM A516", "Grade/Condition": "Grade 55", "Min Yield (ksi)": 30, "Min Yield (MPa)": 205, "Tensile Range (ksi)": "55–75", "Tensile Range (MPa)": "380–515"},
    {"Group": "I", "Specification": "ASTM A516", "Grade/Condition": "Grade 60", "Min Yield (ksi)": 32, "Min Yield (MPa)": 220, "Tensile Range (ksi)": "60–80", "Tensile Range (MPa)": "415–550"},
    {"Group": "I", "Specification": "API 5L", "Grade/Condition": "Grade B", "Min Yield (ksi)": 35, "Min Yield (MPa)": 241, "Tensile Range (ksi)": "60 min.", "Tensile Range (MPa)": "414 min."},
    {"Group": "I", "Specification": "API 5L", "Grade/Condition": "Grade X42", "Min Yield (ksi)": 42, "Min Yield (MPa)": 290, "Tensile Range (ksi)": "60 min.", "Tensile Range (MPa)": "414 min."},
    # GROUP II
    {"Group": "II", "Specification": "ASTM A36", "Grade/Condition": "All thicknesses", "Min Yield (ksi)": 36, "Min Yield (MPa)": 250, "Tensile Range (ksi)": "58–80", "Tensile Range (MPa)": "400–550"},
    {"Group": "II", "Specification": "ASTM A572", "Grade/Condition": "Grade 42", "Min Yield (ksi)": 42, "Min Yield (MPa)": 290, "Tensile Range (ksi)": "60 min.", "Tensile Range (MPa)": "415 min."},
    {"Group": "II", "Specification": "ASTM A572", "Grade/Condition": "Grade 50", "Min Yield (ksi)": 50, "Min Yield (MPa)": 345, "Tensile Range (ksi)": "65 min.", "Tensile Range (MPa)": "450 min."},
    {"Group": "II", "Specification": "ASTM A572", "Grade/Condition": "Grade 55", "Min Yield (ksi)": 55, "Min Yield (MPa)": 380, "Tensile Range (ksi)": "70 min.", "Tensile Range (MPa)": "485 min."},
    {"Group": "II", "Specification": "ASTM A588", "Grade/Condition": "≤4 in [100 mm]", "Min Yield (ksi)": 50, "Min Yield (MPa)": 345, "Tensile Range (ksi)": "70 min.", "Tensile Range (MPa)": "485 min."},
    {"Group": "II", "Specification": "ASTM A709", "Grade/Condition": "Grade 50", "Min Yield (ksi)": 50, "Min Yield (MPa)": 345, "Tensile Range (ksi)": "65 min.", "Tensile Range (MPa)": "450 min."},
    {"Group": "II", "Specification": "ASTM A992", "Grade/Condition": "—", "Min Yield (ksi)": "50–65", "Min Yield (MPa)": "345–450", "Tensile Range (ksi)": "65 min.", "Tensile Range (MPa)": "450 min."},
    {"Group": "II", "Specification": "API 5L", "Grade/Condition": "Grade X52", "Min Yield (ksi)": 52, "Min Yield (MPa)": 359, "Tensile Range (ksi)": "66 min.", "Tensile Range (MPa)": "455 min."},
    # GROUP III
    {"Group": "III", "Specification": "ASTM A572", "Grade/Condition": "Grade 60", "Min Yield (ksi)": 60, "Min Yield (MPa)": 415, "Tensile Range (ksi)": "75 min.", "Tensile Range (MPa)": "520 min."},
    {"Group": "III", "Specification": "ASTM A572", "Grade/Condition": "Grade 65", "Min Yield (ksi)": 65, "Min Yield (MPa)": 450, "Tensile Range (ksi)": "80 min.", "Tensile Range (MPa)": "550 min."},
    {"Group": "III", "Specification": "ASTM A913", "Grade/Condition": "Grade 60", "Min Yield (ksi)": 60, "Min Yield (MPa)": 415, "Tensile Range (ksi)": "75 min.", "Tensile Range (MPa)": "520 min."},
    {"Group": "III", "Specification": "ASTM A913", "Grade/Condition": "Grade 65", "Min Yield (ksi)": 65, "Min Yield (MPa)": 450, "Tensile Range (ksi)": "80 min.", "Tensile Range (MPa)": "550 min."},
    # GROUP IV
    {"Group": "IV", "Specification": "ASTM A709", "Grade/Condition": "Grade HPS70W", "Min Yield (ksi)": 70, "Min Yield (MPa)": 485, "Tensile Range (ksi)": "85–110", "Tensile Range (MPa)": "585–760"},
    {"Group": "IV", "Specification": "ASTM A913", "Grade/Condition": "Grade 70", "Min Yield (ksi)": 70, "Min Yield (MPa)": 485, "Tensile Range (ksi)": "90 min.", "Tensile Range (MPa)": "620 min."},
    {"Group": "IV", "Specification": "ASTM A1066", "Grade/Condition": "Grade 70", "Min Yield (ksi)": 70, "Min Yield (MPa)": 485, "Tensile Range (ksi)": "85 min.", "Tensile Range (MPa)": "585 min."},
    # GROUP V
    {"Group": "V", "Specification": "ASTM A913", "Grade/Condition": "Grade 80", "Min Yield (ksi)": 80, "Min Yield (MPa)": 550, "Tensile Range (ksi)": "95 min.", "Tensile Range (MPa)": "655 min."},
]

# ════════════════════════════════════════════════════════
# NEW: TABLE 5.7 — FILLER METALS (SMAW & SAW)
# ════════════════════════════════════════════════════════
filler_metals = [
    {"Base Metal Group": "I", "Process": "SMAW", "AWS Spec": "A5.1 Carbon Steel", "Electrode Classifications": "F60XX, E70XX"},
    {"Base Metal Group": "I", "Process": "SMAW", "AWS Spec": "A5.5 Low-Alloy Steel", "Electrode Classifications": "E70XX-A1, E70XX-C1L, E70XX-C2L, E70XX-C3L, E7010-P1, E7018-W1"},
    {"Base Metal Group": "I", "Process": "SAW", "AWS Spec": "A5.17 Carbon Steel", "Electrode Classifications": "F6XX-EXXX, F6XX-ECXXX, F7XX-EXXX, F7XX-ECXXX"},
    {"Base Metal Group": "I", "Process": "SAW", "AWS Spec": "A5.23 Low-Alloy Steel", "Electrode Classifications": "F7XX-E(C)XXX-A1, F7XX-E(C)XXX-A2, F7XX-E(C)XXX-A4, F7XX-E(C)XXX-Ni1, F7XX-E(C)XXX-Ni2"},
    {"Base Metal Group": "II", "Process": "SMAW", "AWS Spec": "A5.1 Carbon Steel", "Electrode Classifications": "E7015, E7016, E7018, E7028"},
    {"Base Metal Group": "II", "Process": "SMAW", "AWS Spec": "A5.5 Low-Alloy Steel", "Electrode Classifications": "E7015-A1, E7016-A1, E7018-A1, E7015-C2L, E7016-C2L, E7018-C2L, E7018-C3L, E7018-W1"},
    {"Base Metal Group": "II", "Process": "SAW", "AWS Spec": "A5.17 Carbon Steel", "Electrode Classifications": "F7XX-EXXX, F7XX-ECXXX"},
    {"Base Metal Group": "II", "Process": "SAW", "AWS Spec": "A5.23 Low-Alloy Steel", "Electrode Classifications": "F7XX-E(C)XXX-A1, F7XX-E(C)XXX-A2, F7XX-E(C)XXX-A4, F7XX-E(C)XXX-Ni1, F7XX-E(C)XXX-Ni2"},
    {"Base Metal Group": "III", "Process": "SMAW", "AWS Spec": "A5.5 Low-Alloy Steel", "Electrode Classifications": "E8016-C1, E8018-C1, E8016-C2, E8018-C2, E8016-C3, E8018-C3, E8016-C4, E8018-C4, E8018-MN1, E8018-W2"},
    {"Base Metal Group": "III", "Process": "SAW", "AWS Spec": "A5.23 Low-Alloy Steel", "Electrode Classifications": "F8AX-E(C)XXX-A1 thru A4, F8AX-E(C)XXX-Ni1 thru Ni5, F8AX-E(C)XXX-W"},
    {"Base Metal Group": "IV", "Process": "SMAW", "AWS Spec": "A5.5 Low-Alloy Steel", "Electrode Classifications": "E9018M, E9018-P2"},
    {"Base Metal Group": "IV", "Process": "SAW", "AWS Spec": "A5.23 Low-Alloy Steel", "Electrode Classifications": "F9AX-E(C)XXX-A1, A2, A3, F3, M2, Ni4"},
    {"Base Metal Group": "V", "Process": "SMAW", "AWS Spec": "A5.5 Low-Alloy Steel", "Electrode Classifications": "E10018M, E10045-P2"},
    {"Base Metal Group": "V", "Process": "SAW", "AWS Spec": "A5.23 Low-Alloy Steel", "Electrode Classifications": "F10AX-EXX-XX, F10AX-ECXX-XX"},
]

# ════════════════════════════════════════════════════════
# NEW: TABLE 5.10 — SHIELDING GASES FOR GMAW
# ════════════════════════════════════════════════════════
shielding_gas_data = [
    {"Electrode": "ER70S-X (except ER70S-G) and E70C-X metal cored", "Shielding Gas Type": "Ar/CO₂ combinations", "Composition": "Ar 75–90% / CO₂ 10–25%"},
    {"Electrode": "ER70S-X (except ER70S-G) and E70C-X metal cored", "Shielding Gas Type": "Ar/O₂ combinations", "Composition": "Ar 95–98% / O₂ 2–5%"},
    {"Electrode": "ER70S-X (except ER70S-G) and E70C-X metal cored", "Shielding Gas Type": "100% CO₂", "Composition": "100% CO₂"},
]

# ════════════════════════════════════════════════════════
# UI — TABS (now 8 tabs)
# ════════════════════════════════════════════════════════
tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10 = st.tabs([
    "⚡ SMAW (5.1)",
    "🔩 SAW (5.2)",
    "🌀 GMAW (5.3)",
    "🔥 FCAW (5.4)",
    "🌡️ Preheat (5.11)",
    "📐 PJP Min Size (5.8)",
    "🏗️ Base Metals (5.6)",
    "🔧 Filler Metals (5.7)",
    "🤖 AI Assistant",
    "📄 WPS Generator",
])

def position_filter(df, pos):
    if pos == "All":
        return df
    return df[df["Position"].str.contains(pos, case=False, na=False)]

# ── TAB 1: SMAW ──
with tab1:
    st.subheader("Table 5.1 — Prequalified SMAW WPS Requirements")
    st.caption("Shielded Metal Arc Welding")
    pos = st.selectbox("Filter by Position:", ["All", "Flat", "Horizontal", "Vertical", "Overhead"], key="smaw_pos")
    df = pd.DataFrame(smaw_data)
    st.dataframe(position_filter(df, pos), use_container_width=True, hide_index=True)

    st.divider()
    st.subheader("📐 Joint Illustration — Figure 5.1")

    from PIL import Image
    import os

    image_path = "images/figure_5_1.png"

    if os.path.exists(image_path):
        image = Image.open(image_path)
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.image(image, caption="Figure 5.1 — Prequalified CJP Groove Welded Joint Details — AWS D1.1/D1.1M:2025", width=500)
    else:
        st.warning("⚠️ Image not found. Please add figure_5_1.png to the images folder.")

        
# ── TAB 2: SAW ──
with tab2:
    st.subheader("Table 5.2 — Prequalified SAW WPS Requirements")
    st.caption("Submerged Arc Welding")
    pos = st.selectbox("Filter by Position:", ["All", "Flat", "Horizontal"], key="saw_pos")
    df = pd.DataFrame(saw_data)
    st.dataframe(position_filter(df, pos), use_container_width=True, hide_index=True)
    st.info("⚠️ SAW is only prequalified for Flat and Horizontal positions.")

# ── TAB 3: GMAW ──
with tab3:
    st.subheader("Table 5.3 — Prequalified GMAW (Solid Wire) WPS Requirements")
    st.caption("Gas Metal Arc Welding — solid wire only. GMAW-S is NOT prequalified.")
    pos = st.selectbox("Filter by Position:", ["All", "Flat", "Horizontal", "Vertical", "Overhead"], key="gmaw_pos")
    df = pd.DataFrame(gmaw_data)
    st.dataframe(position_filter(df, pos), use_container_width=True, hide_index=True)
    st.warning("⚠️ GMAW-S (short circuit transfer) is NOT prequalified. Must use CV power source.")

# ── TAB 4: FCAW ──
with tab4:
    st.subheader("Table 5.4 — Prequalified FCAW & GMAW Metal Cored WPS Requirements")
    pos = st.selectbox("Filter by Position:", ["All", "Flat", "Horizontal", "Vertical", "Overhead"], key="fcaw_pos")
    df = pd.DataFrame(fcaw_data)
    st.dataframe(position_filter(df, pos), use_container_width=True, hide_index=True)

# ── TAB 5: PREHEAT ──
with tab5:
    st.subheader("Table 5.11 — Minimum Preheat & Interpass Temperature")
    col1, col2 = st.columns(2)
    with col1:
        cat = st.selectbox("Filter by Category:", ["All", "A", "B", "C"], key="preheat_cat")
    with col2:
        unit = st.radio("Temperature Unit:", ["°F", "°C"], horizontal=True)
    df = pd.DataFrame(preheat_data)
    if cat != "All":
        df = df[df["Category"] == cat]
    if unit == "°F":
        df = df.drop(columns=["Min Preheat °C"])
    else:
        df = df.drop(columns=["Min Preheat °F"])
    st.dataframe(df, use_container_width=True, hide_index=True)
    with st.expander("📋 Category Steel Specifications"):
        st.markdown("""
**Category A** — A36 (≤3/4 in), A53 Gr.B, A106 Gr.B, A131, A139, A381, A500, A501, A516 Gr.55/60, A524, A573, A709 Gr.36, A1008/A1011/A1018 SS, API 5L Gr.B & X42, ABS Gr.A,B,D,E

**Category B** — A36 (all), A131 AH/DH/EH 32/36, A216, A501 Gr.B, A516 Gr.65/70, A529, A537 Cl.1/2, A572 Gr.42/50/55, A588, A595, A606, A618, A633 Gr.A/C/D, A709 Gr.36/50/50W/50S/HPS50W, A710, A847, A913 Gr.50/60/65, A992, A1008/A1011/A1018 HSLAS, API 2H/2MT1/2W/2Y Gr.42/50, API 5L X52, ABS AH/DH/EH 32/36

**Category C** — A572 Gr.60/65, A633 Gr.E, A709 HPS70W, A710 (Class 2/3), A913 Gr.70, A1018 HSLAS Gr.60/70 Cl.2, A1066 Gr.60/65/70, API 2W/2Y Gr.60, API 5L X52
        """)

# ── TAB 6: PJP MIN SIZE ──
with tab6:
    st.subheader("Table 5.8 — Minimum Prequalified PJP Groove Weld Size")
    st.caption("Per 5.4.2.3(1) — weld size need not exceed thickness of the thinner part joined")
    unit = st.radio("Display units:", ["Both", "Inches only", "mm only"], horizontal=True, key="pjp_unit")
    df = pd.DataFrame(pjp_size_data)
    if unit == "Inches only":
        df = df[["Base Metal Thickness (in)", "Min Weld Size (in)"]]
    elif unit == "mm only":
        df = df[["Base Metal Thickness (mm)", "Min Weld Size (mm)"]]
    st.dataframe(df, use_container_width=True, hide_index=True)

# ── TAB 7: BASE METALS ──
with tab7:
    st.subheader("Table 5.6 — Approved Base Metals for Prequalified WPSs")
    st.caption("Only these base metals may be used in prequalified WPSs (see 5.3)")

    col1, col2 = st.columns(2)
    with col1:
        group_filter = st.selectbox("Filter by Group:", ["All", "I", "II", "III", "IV", "V"])
    with col2:
        search = st.text_input("Search specification (e.g. A572, A36, API):", "")

    df = pd.DataFrame(base_metals)
    if group_filter != "All":
        df = df[df["Group"] == group_filter]
    if search:
        df = df[df["Specification"].str.contains(search, case=False, na=False)]

    st.dataframe(df, use_container_width=True, hide_index=True)

    with st.expander("ℹ️ Group Notes"):
        st.markdown("""
- **Group I** — Lower strength steels (Fy ≤ 36 ksi / 250 MPa typically)
- **Group II** — Medium strength steels (Fy up to ~55 ksi / 380 MPa)
- **Group III** — Higher strength steels (Fy 60–65 ksi / 415–450 MPa)
- **Group IV** — High strength steels (Fy ~70 ksi / 485 MPa)
- **Group V** — Very high strength steels (Fy ~80 ksi / 550 MPa)

**Note:** In joints with base metals of different groups, use filler metal matching the higher strength base metal, or use filler matching the lower strength base metal if it produces a low-hydrogen deposit.
        """)

# ── TAB 8: FILLER METALS ──
with tab8:
    st.subheader("Table 5.7 — Filler Metals for Matching Strength")
    st.caption("SMAW and SAW filler metals for Table 5.6 Groups I–V")

    col1, col2 = st.columns(2)
    with col1:
        grp = st.selectbox("Base Metal Group:", ["All", "I", "II", "III", "IV", "V"], key="fm_group")
    with col2:
        proc = st.selectbox("Process:", ["All", "SMAW", "SAW"], key="fm_proc")

    df = pd.DataFrame(filler_metals)
    if grp != "All":
        df = df[df["Base Metal Group"] == grp]
    if proc != "All":
        df = df[df["Process"] == proc]
    st.dataframe(df, use_container_width=True, hide_index=True)

    st.divider()
    st.subheader("Table 5.10 — Shielding Gas Options for GMAW (AWS A5.18/A5.18M)")
    st.dataframe(pd.DataFrame(shielding_gas_data), use_container_width=True, hide_index=True)

    with st.expander("📋 Weathering Steel Requirements (Table 5.9 — see 5.6.3)"):
        st.markdown("""
For **exposed bare unpainted weathering steel** applications:

| Process | AWS Spec | Approved Electrodes |
|---|---|---|
| SMAW | A5.5/A5.5M | All electrodes depositing weld metal with B2L, C1, C1L, C2, C2L, C3, or WX analysis |
| SAW | A5.23/A5.23M | All electrode-flux combinations depositing weld metal with Ni1, Ni2, Ni3, Ni4, or WX analysis |
| FCAW | A5.29/A5.29M | All electrodes depositing weld metal with B2L, K2, Ni1, Ni2, Ni3, Ni4, or WX analysis |
| GMAW | A5.28/A5.28M | All electrodes meeting B2L, G, Ni1, Ni2, Ni3 analysis requirements |

**Exceptions (5.6.3.1 & 5.6.3.2):** Single-pass groove welds and small single-pass fillet welds (SMAW ≤ 1/4 in, SAW/GMAW/FCAW ≤ 5/16 in) may use any Group II filler metal from Table 5.7.
        """)


# ── Standards Index (sidebar) ────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### 📚 Standards Index")
    st.caption("Upload PDFs to enable clause-level AI answers.")

    try:
        from rag_engine import WeldingRAG
        rag = WeldingRAG()

        count = rag.index_count()
        if count:
            st.success(f"✅ {count:,} clauses indexed")
        else:
            st.warning("No standards indexed yet.")

        uploaded_pdfs = st.file_uploader(
            "Upload standard PDFs",
            type=["pdf", "txt"],
            accept_multiple_files=True,
            help="AWS D1.1, EN ISO 15614, EN 1011, etc. Never committed to git.",
        )

        if uploaded_pdfs and st.button("🔍 Index uploaded PDFs", use_container_width=True):
            import tempfile, os
            os.makedirs("standards", exist_ok=True)
            saved_paths = []
            for f in uploaded_pdfs:
                tmp_path = f"standards/{f.name}"
                with open(tmp_path, "wb") as fp:
                    fp.write(f.read())
                saved_paths.append(tmp_path)

            log = st.empty()
            messages = []
            def progress(msg):
                messages.append(msg)
                log.markdown("\\n".join(messages))

            with st.spinner("Indexing…"):
                added = rag.build_index(saved_paths, progress_cb=progress)
            st.success(f"Done — {added} new chunks added.")
            st.rerun()

        if count and st.button("🗑️ Clear index", use_container_width=True):
            rag.build_index([], reset=True)
            st.rerun()

    except ImportError:
        st.info("Install RAG packages to enable clause retrieval:\\n`pip install pymupdf chromadb sentence-transformers`")

with tab9:
    st.subheader("🤖 AI Welding Assistant")
    st.markdown(
        "Fill in your welding parameters, then chat with the AI. "
        "When standard PDFs are indexed, answers are grounded in real clauses."
    )
    st.divider()

    # ── RAG availability check ────────────────────────────────────────────────
    try:
        from rag_engine import WeldingRAG
        _rag = WeldingRAG()
        RAG_READY = _rag.is_ready()
    except Exception:
        _rag      = None
        RAG_READY = False

    if RAG_READY:
        st.success(f"📚 Standards index active — {_rag.index_count():,} clauses loaded.", icon="✅")
    else:
        st.info(
            "💡 No standards indexed yet. Answers use hardcoded table data only. "
            "Upload PDFs in the sidebar to unlock full clause retrieval.",
            icon="ℹ️",
        )

    st.divider()

    # ── Parameter inputs ──────────────────────────────────────────────────────
    col1, col2 = st.columns(2)
    with col1:
        process   = st.selectbox("Welding Process:",  ["SMAW", "SAW", "GMAW", "FCAW"])
        position  = st.selectbox("Welding Position:", ["Flat", "Horizontal", "Vertical", "Overhead"])
        weld_type = st.selectbox("Weld Type:",        ["Fillet", "Groove", "Root Pass"])
    with col2:
        base_metal = st.selectbox("Base Metal:", [
            "ASTM A36", "ASTM A572 Gr.50", "ASTM A572 Gr.60",
            "ASTM A588", "ASTM A709 Gr.50", "ASTM A992",
            "API 5L Gr.B", "Other"
        ])
        thickness = st.number_input("Material Thickness (mm):", min_value=1.0, max_value=200.0, value=12.0, step=0.5)
        electrode = st.text_input("Electrode/Wire (optional):", placeholder="e.g. E7018, ER70S-6")

    col3, col4 = st.columns(2)
    with col3:
        ambient_temp = st.number_input("Ambient Temperature (°C):", min_value=-50.0, max_value=50.0, value=20.0, step=1.0)
        humidity     = st.selectbox("Humidity/Weather Condition:", [
            "Normal (indoor)", "High humidity", "Rain/Snow nearby",
            "Strong wind", "Extreme cold (below 0°C)", "Extreme heat (above 35°C)"
        ])
    with col4:
        location    = st.selectbox("Welding Location:", [
            "Indoor — controlled environment", "Outdoor — sheltered",
            "Outdoor — exposed", "Offshore / marine environment", "Underground"
        ])
        criticality = st.selectbox("Structure Criticality:", [
            "Standard structure",
            "Critical structure (bridges, pressure vessels)",
            "Seismic zone",
            "Impact/dynamic loading"
        ])

    additional = st.text_area(
        "Describe your full situation in detail:",
        placeholder="e.g. Welding a structural beam connection outdoors in winter. -20°C, wind present. Bridge structure…",
        height=100
    )
    st.divider()

    # ── AWS D1.1 hardcoded reference data (unchanged) ─────────────────────────
    smaw_rules = {
        "Flat":       {"Fillet":    {"max_electrode": "5/16 in [8.0 mm]", "exception": "Except root pass"},
                       "Groove":    {"max_electrode": "1/4 in [6.4 mm]",  "exception": "Except root pass"},
                       "Root Pass": {"max_electrode": "3/16 in [4.8 mm]", "exception": ""}},
        "Horizontal": {"Fillet":    {"max_electrode": "1/4 in [6.4 mm]",  "exception": ""},
                       "Groove":    {"max_electrode": "3/16 in [4.8 mm]", "exception": ""}},
        "Vertical":   {"All":       {"max_electrode": "3/16 in [4.8 mm]", "exception": "5/32 in [4.0 mm] for EXX14 and low-hydrogen electrodes"}},
        "Overhead":   {"All":       {"max_electrode": "3/16 in [4.8 mm]", "exception": "5/32 in [4.0 mm] for EXX14 and low-hydrogen electrodes"}},
    }
    saw_rules = {
        "max_electrode": "1/4 in [6.4 mm]", "positions": "Flat and Horizontal only",
        "current_fillet": "1000A max (single), 1200A max (parallel), Unlimited (multiple)",
        "current_groove_root_with_opening": "600A max (single), 700A max (parallel)",
        "current_groove_root_no_opening": "900A max (single)",
        "current_groove_fill": "1200A max (single)",
    }
    gmaw_rules = {
        "max_electrode": "0.0625 in [1.6 mm]",
        "note": "GMAW-S is NOT prequalified. Must use CV power source.",
        "min_current": "190A for 0.030in | 210A for 0.035in | 230A for 0.040in | 260A for 0.045in | 300A for 0.0625in",
    }
    fcaw_rules = {
        "flat_horizontal_T_less_10mm": "0.030 in [0.8 mm] min (FCAW), No min (GMAW Cored)",
        "flat_horizontal_T_more_10mm": "0.045 in [1 mm] min, 1/8 in [3.2 mm] max (FCAW)",
        "vertical_T_more_10mm": "3/32 in [2.4 mm] max",
        "overhead_T_more_10mm": "5/64 in [2.2 mm] max",
    }
    bead_limits = {
        "SMAW": {
            "root_bead":          {"Flat": "3/8 in [10 mm]", "Horizontal": "5/16 in [8 mm]", "Vertical": "1/2 in [12 mm]", "Overhead": "5/16 in [8 mm]"},
            "fill_cap":           "3/16 in [5 mm] all positions",
            "single_pass_fillet": {"Flat": "3/8 in [10 mm]", "Horizontal": "5/16 in [8 mm]", "Vertical": "1/2 in [12 mm]", "Overhead": "5/16 in [8 mm]"},
        },
        "SAW": {
            "root_bead":          "Unlimited (Flat & Horizontal only)",
            "fill_cap_single":    "1/4 in [6 mm] max",
            "fill_cap_multiple":  "Unlimited",
            "single_pass_fillet": {"single": "5/16 in [8 mm]", "parallel": "5/16 in [8 mm]", "multiple": "1/2 in [12 mm]"},
        },
        "GMAW": {
            "root_bead":          {"Flat": "3/8 in [10 mm]", "Horizontal": "5/16 in [8 mm]", "Vertical": "1/2 in [12 mm]", "Overhead": "5/16 in [8 mm]"},
            "fill_cap":           "1/4 in [6 mm] all positions",
            "single_pass_fillet": {"Flat": "1/2 in [12 mm]", "Horizontal": "3/8 in [10 mm]", "Vertical": "1/2 in [12 mm]", "Overhead": "5/16 in [8 mm]"},
        },
        "FCAW": {
            "root_bead":          {"Flat": "3/8 in [10 mm]", "Horizontal": "5/16 in [8 mm]", "Vertical": "1/2 in [12 mm]", "Overhead": "5/16 in [8 mm]"},
            "fill_cap":           "1/4 in [6 mm] all positions",
            "single_pass_fillet": {"Flat": "1/2 in [12 mm]", "Horizontal": "3/8 in [10 mm]", "Vertical": "1/2 in [12 mm]", "Overhead": "5/16 in [8 mm]"},
        },
    }
    preheat_rules = {
        "ASTM A36": {
            "category": "A (SMAW non-LH) or B (SMAW LH, SAW, GMAW, FCAW)",
            "Cat_A": {"up_to_19mm": "32°F [0°C]",  "19_to_38mm": "150°F [65°C]",  "38_to_65mm": "225°F [110°C]", "over_65mm": "300°F [150°C]"},
            "Cat_B": {"up_to_19mm": "32°F [0°C]",  "19_to_38mm": "50°F [10°C]",   "38_to_65mm": "150°F [65°C]",  "over_65mm": "225°F [110°C]"},
        },
        "ASTM A572 Gr.50": {
            "category": "B",
            "Cat_B": {"up_to_19mm": "32°F [0°C]",  "19_to_38mm": "50°F [10°C]",   "38_to_65mm": "150°F [65°C]",  "over_65mm": "225°F [110°C]"},
        },
        "ASTM A572 Gr.60": {
            "category": "C",
            "Cat_C": {"up_to_19mm": "50°F [10°C]", "19_to_38mm": "150°F [65°C]",  "38_to_65mm": "225°F [110°C]", "over_65mm": "300°F [150°C]"},
        },
    }

    # ── Resolve table data for current inputs ─────────────────────────────────
    if process == "SMAW":
        wt   = weld_type if weld_type in smaw_rules.get(position, {}) else "All"
        rule = smaw_rules.get(position, {}).get(wt, smaw_rules.get(position, {}).get("All", {}))
        process_data = (
            f"Max electrode diameter: {rule.get('max_electrode', 'N/A')}\n"
            f"Exception: {rule.get('exception', 'None')}"
        )
    elif process == "SAW":
        r = saw_rules
        process_data = (
            f"Max electrode diameter: {r['max_electrode']}\n"
            f"Approved positions: {r['positions']}\n"
            f"Current — fillet: {r['current_fillet']}\n"
            f"Current — groove root (with opening): {r['current_groove_root_with_opening']}\n"
            f"Current — groove root (no opening): {r['current_groove_root_no_opening']}\n"
            f"Current — groove fill: {r['current_groove_fill']}"
        )
    elif process == "GMAW":
        r = gmaw_rules
        process_data = (
            f"Max electrode diameter: {r['max_electrode']}\n"
            f"Note: {r['note']}\n"
            f"Minimum current by wire size: {r['min_current']}"
        )
    elif process == "FCAW":
        r = fcaw_rules
        process_data = (
            f"Flat/Horizontal, t < 10 mm: {r['flat_horizontal_T_less_10mm']}\n"
            f"Flat/Horizontal, t ≥ 10 mm: {r['flat_horizontal_T_more_10mm']}\n"
            f"Vertical, t ≥ 10 mm: {r['vertical_T_more_10mm']}\n"
            f"Overhead, t ≥ 10 mm: {r['overhead_T_more_10mm']}"
        )

    bl = bead_limits.get(process, {})
    bead_data_lines = []
    if "root_bead" in bl:
        rb = bl["root_bead"]
        bead_data_lines.append(f"Root bead max: {rb[position] if isinstance(rb, dict) else rb}")
    if "fill_cap" in bl:
        bead_data_lines.append(f"Fill/cap bead max: {bl['fill_cap']}")
    if "fill_cap_single" in bl:
        bead_data_lines.append(f"Fill/cap single-wire max: {bl['fill_cap_single']}")
        bead_data_lines.append(f"Fill/cap multi-wire max:  {bl['fill_cap_multiple']}")
    if "single_pass_fillet" in bl:
        spf = bl["single_pass_fillet"]
        bead_data_lines.append(f"Single-pass fillet max: {spf[position] if isinstance(spf, dict) else spf.get('single', spf)}")
    bead_data = "\n".join(bead_data_lines)

    thickness_cat = (
        "up_to_19mm"  if thickness <= 19 else
        "19_to_38mm"  if thickness <= 38 else
        "38_to_65mm"  if thickness <= 65 else
        "over_65mm"
    )
    ph_metal = preheat_rules.get(base_metal)
    if ph_metal:
        if "Cat_B" in ph_metal:   cat_key = "Cat_B"
        elif "Cat_A" in ph_metal: cat_key = "Cat_A"
        else:                     cat_key = list(ph_metal.keys())[-1]
        ph_temp      = ph_metal[cat_key].get(thickness_cat, "N/A")
        preheat_data = (
            f"Steel category: {ph_metal['category']}\n"
            f"Min preheat for {thickness_cat.replace('_',' ')}: {ph_temp}"
        )
    else:
        preheat_data = "Refer to Table 5.11 — base metal not in prequalified list."

    situation_block = f"""
WELDING JOB PARAMETERS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Process          : {process}
Position         : {position}
Weld type        : {weld_type}
Base metal       : {base_metal}
Thickness        : {thickness} mm  ({thickness_cat.replace('_',' ')})
Electrode / wire : {electrode if electrode else 'not specified'}
Ambient temp     : {ambient_temp}°C
Weather          : {humidity}
Location         : {location}
Criticality      : {criticality}
Additional notes : {additional if additional else 'none'}

AWS D1.1 TABLE DATA (hardcoded)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PROCESS LIMITS (Tables 5.1–5.4):
{process_data}

BEAD THICKNESS LIMITS (Tables 5.1–5.4):
{bead_data}

PREHEAT (Table 5.11):
{preheat_data}
""".strip()

    # ── Build system prompt (RAG-aware) ───────────────────────────────────────
    rag_instruction = (
        "RETRIEVED STANDARD CLAUSES are provided above the user's question. "
        "These are verbatim excerpts from the actual standards. "
        "Prioritise these clauses over your general training knowledge. "
        "Always cite the source name when referencing a retrieved clause."
        if RAG_READY else
        "No standard PDFs are indexed yet. Use only the hardcoded table data "
        "provided in the job parameters. If a question goes beyond that data, "
        "say so clearly and cite the correct clause reference for the user to check."
    )

    SYSTEM_PROMPT = f"""You are a senior International Welding Engineer (IWE) with 20 years of hands-on \
experience in structural steel fabrication. You have deep expertise in AWS D1.1/D1.1M:2025, \
EN ISO 15614, EN 1011, and practical shop-floor and field welding.

YOUR BEHAVIOUR RULES:
1. Cite exact clause numbers or table references for every requirement you state.
2. When table data or retrieved clauses are provided, use those exact values. Never invent numbers.
3. When data is absent, say so clearly and give the correct clause reference to consult.
4. Flag safety-critical issues with ⚠️ and explain WHY they matter, not just WHAT they are. \
For SMAW specifically: wind does NOT disrupt shielding gas (SMAW has none) — \
wind accelerates heat loss from preheated steel, increasing hydrogen cracking risk. \
Never confuse SMAW wind risk with GMAW/FCAW shielding gas disruption.
5. Give practical advice — not just "comply with the standard" but HOW to comply on the job.
6. If a question is outside AWS D1.1 scope (e.g. EN ISO, API), answer from your IWE knowledge \
and label it clearly as outside D1.1 scope.
7. Be direct and opinionated. If conditions are dangerous, say so firmly.
8. Write like an experienced engineer advising a colleague — structured but not bureaucratic.
9. Before giving a Go/No-Go recommendation, explicitly state: \
the minimum preheat required, the current ambient temperature, \
the difference between them, and whether welding can start RIGHT NOW \
or only after preheating. Never give a Conditional GO if ambient \
temperature is below the minimum preheat — that is a NO-GO until \
the steel is physically heated to the minimum temperature.

REGARDING RETRIEVED CLAUSES:
{rag_instruction}

THE WELDING JOB YOU ARE ADVISING ON:
{situation_block}

The user will now ask questions about this job. Answer specifically in context of the parameters above.
"""

    # ── Session state ─────────────────────────────────────────────────────────
    param_key = (process, position, weld_type, base_metal, thickness,
                 electrode, ambient_temp, humidity, location, criticality, additional, RAG_READY)

    if "ai_chat_history" not in st.session_state:
        st.session_state.ai_chat_history = []
        st.session_state.ai_param_key    = param_key

    if st.session_state.ai_param_key != param_key:
        st.session_state.ai_chat_history = []
        st.session_state.ai_param_key    = param_key
        st.info("⚙️ Parameters changed — conversation reset.")

    # ── Quick-start buttons ───────────────────────────────────────────────────
    st.markdown("**Quick analysis — or type your own question below:**")
    qcol1, qcol2, qcol3 = st.columns(3)
    quick_prompt = None
    if qcol1.button("📋 Full job analysis",       use_container_width=True):
        quick_prompt = (
            "Give me a full analysis of this welding job: electrode limits, bead thickness limits, "
            "preheat requirements, environmental risks, and a go/no-go recommendation."
        )
    if qcol2.button("🌡️ Preheat & hydrogen risk", use_container_width=True):
        quick_prompt = (
            "Focus on preheat and hydrogen cracking risk for this job. "
            "Is the ambient temperature acceptable? What precautions do I need?"
        )
    if qcol3.button("⚠️ Environmental risks",     use_container_width=True):
        quick_prompt = (
            "What are the main environmental risks for this welding job "
            "and what specific precautions should I take on site?"
        )

    # ── Chat display ──────────────────────────────────────────────────────────
    for msg in st.session_state.ai_chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # ── Input ─────────────────────────────────────────────────────────────────
    user_input = st.chat_input("Ask anything about this welding job…")
    if quick_prompt and not user_input:
        user_input = quick_prompt

    if user_input:
        st.session_state.ai_chat_history.append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)

        # ── RAG retrieval ─────────────────────────────────────────────────────
        rag_context = ""
        if RAG_READY and _rag:
            # Build a rich search query from the user question + job parameters
            search_query = (
                f"{user_input} "
                f"{process} {position} {weld_type} {base_metal} "
                f"{thickness}mm preheat electrode bead"
            )
            try:
                chunks      = _rag.query(search_query, n=7)
                rag_context = _rag.format_context(chunks) if chunks else ""
            except Exception as e:
                rag_context = f"(RAG retrieval failed: {e})"

        # Build messages array: system + trimmed history
        # Keep last 10 messages max to avoid context window errors
        MAX_HISTORY = 10
        trimmed_history = st.session_state.ai_chat_history[-MAX_HISTORY:]

        messages_for_api = [{"role": "system", "content": SYSTEM_PROMPT}]

        if rag_context:
            last_user = trimmed_history[-1]
            prior     = trimmed_history[:-1]
            for m in prior:
                messages_for_api.append({"role": m["role"], "content": m["content"]})
            messages_for_api.append({
                "role":    "system",
                "content": f"RETRIEVED STANDARD CLAUSES FOR THIS QUESTION:\n\n{rag_context}",
            })
            messages_for_api.append({"role": "user", "content": last_user["content"]})
        else:
            for m in trimmed_history:
                messages_for_api.append({"role": m["role"], "content": m["content"]})

        # Inject retrieved clauses as a system-level context message
        # (placed just before the latest user message for maximum attention)
        history_to_send = list(st.session_state.ai_chat_history)
        if rag_context:
            # Insert retrieved context right before the last user message
            last_user = history_to_send.pop()
            messages_for_api += [{"role": m["role"], "content": m["content"]}
                                  for m in history_to_send[:-0] or history_to_send]
            messages_for_api.append({
                "role":    "system",
                "content": f"RETRIEVED STANDARD CLAUSES FOR THIS QUESTION:\n\n{rag_context}",
            })
            messages_for_api.append({"role": "user", "content": last_user["content"]})
        else:
            messages_for_api += [{"role": m["role"], "content": m["content"]}
                                  for m in history_to_send]

        # ── API call ──────────────────────────────────────────────────────────
        with st.chat_message("assistant"):
            with st.spinner("Analysing…"):
                try:
                    from groq import Groq
                    client   = Groq(api_key=st.secrets["GROQ_API_KEY"])
                    response = client.chat.completions.create(
                        model       = "llama-3.3-70b-versatile",
                        messages    = messages_for_api,
                        max_tokens  = 2000,
                        temperature = 0.15,
                    )
                    reply = response.choices[0].message.content
                    st.markdown(reply)
                    st.session_state.ai_chat_history.append({"role": "assistant", "content": reply})

                    # Show which clauses were used (expandable, not intrusive)
                    if rag_context and chunks:
                        with st.expander(f"📖 {len(chunks)} standard clauses retrieved", expanded=False):
                            for i, c in enumerate(chunks, 1):
                                st.caption(f"**[{i}] {c['source']}** — relevance {1-c['distance']:.0%}")
                                st.text(c["text"][:400] + ("…" if len(c["text"]) > 400 else ""))

                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")

    # ── Reset button ──────────────────────────────────────────────────────────
    if st.session_state.ai_chat_history:
        st.divider()
        if st.button("🗑️ Clear conversation", use_container_width=False):
            st.session_state.ai_chat_history = []
            st.rerun()

    st.divider()
    st.caption("⚠️ AI interpretation is for engineering reference only. Always verify against the official AWS D1.1/D1.1M:2025 standard before production welding.")




with tab10:
    st.subheader("📄 AWS D1.1/D1.1M:2020 Form J-2 (GMAW & FCAW) Generator")
    st.markdown("Generate a professional WPS document matching the **official AWS D1.1/D1.1M:2020 Annex J Form J-2** layout for GMAW/FCAW processes.")
    st.divider()

    # ── SECTION 1: HEADER INFO ──
    st.markdown("### 🏢 Header Information")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        company_name = st.text_input("Company Name:", placeholder="e.g. Genova Steel Ltd")
    with col2:
        wps_number = st.text_input("WPS No.:", placeholder="e.g. WPS-001-2025")
    with col3:
        rev_number = st.text_input("Rev. No.:", value="0")
    with col4:
        wps_date = st.date_input("Date:")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        authorized_by = st.text_input("Authorized By:", placeholder="e.g. Babak Pirzadi")
    with col2:
        auth_date = st.date_input("Authorization Date:")
    with col3:
        supporting_pqr = st.text_input("Supporting PQR(s):", placeholder="e.g. PQR-001-2025")
    with col4:
        cvn_report = st.text_input("CVN Report:", placeholder="e.g. N/A")

    st.divider()

    # ── SECTION 2: BASE METALS ──
    st.markdown("### 🏗️ Base Metals")
    col1, col2 = st.columns(2)
    with col1:
        bm_spec = st.text_input("Base Material Specification:", placeholder="e.g. ASTM A36")
        bm_type_grade = st.text_input("Type or Grade:", placeholder="e.g. Grade B")
        bm_aws_group = st.text_input("AWS Group No.:", placeholder="e.g. Group I")
        bm_welded_to = st.text_input("Welded To (Specification):", placeholder="e.g. ASTM A572 Gr.50")
        bm_backing = st.text_input("Backing Material:", placeholder="e.g. Steel / None")
        bm_diameter = st.text_input("Diameter (pipe/tube):", placeholder="e.g. 168.3 mm OD / N/A")

    with col2:
        st.markdown("**Base Metal Thickness Ranges:**")
        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown("*As-Welded*")
            thick_cjp_aw = st.text_input("CJP Groove Welds:", placeholder="e.g. 6-25 mm", key="cjp_aw")
            thick_cjp_cvn_aw = st.text_input("CJP Groove w/CVN:", placeholder="e.g. N/A", key="cjp_cvn_aw")
            thick_pjp_aw = st.text_input("PJP Groove Welds:", placeholder="e.g. 6-50 mm", key="pjp_aw")
            thick_fillet_aw = st.text_input("Fillet Welds:", placeholder="e.g. All", key="fillet_aw")
        with col_b:
            st.markdown("*With PWHT*")
            thick_cjp_pwht = st.text_input("CJP Groove Welds:", placeholder="e.g. N/A", key="cjp_pwht")
            thick_cjp_cvn_pwht = st.text_input("CJP Groove w/CVN:", placeholder="e.g. N/A", key="cjp_cvn_pwht")
            thick_pjp_pwht = st.text_input("PJP Groove Welds:", placeholder="e.g. N/A", key="pjp_pwht")
            thick_fillet_pwht = st.text_input("Fillet Welds:", placeholder="e.g. N/A", key="fillet_pwht")

    st.divider()

    # ── SECTION 3: JOINT DETAILS ──
    st.markdown("### 🔧 Joint Details")
    col1, col2 = st.columns(2)
    with col1:
        groove_type = st.selectbox("Groove Type:",
            ["Single-V", "Double-V", "Single-Bevel", "Double-Bevel",
             "Single-U", "Double-U", "Square", "Fillet"])
        groove_angle = st.text_input("Groove Angle:", placeholder="e.g. 60°")
        root_opening = st.text_input("Root Opening:", placeholder="e.g. 0-3 mm")
        root_face = st.text_input("Root Face:", placeholder="e.g. 0-3 mm")
        backgouging = st.selectbox("Backgouging:", ["Yes", "No"])
        backgouging_method = st.text_input("Backgouging Method:", placeholder="e.g. Carbon arc / Grinding / N/A")
    with col2:
        st.markdown("**Joint Sketch Description:**")
        joint_sketch = st.text_area("Joint Sketch Notes:",
            placeholder="Describe the joint geometry or reference drawing number...",
            height=180)

    st.divider()

    # ── SECTION 4: POSTWELD HEAT TREATMENT ──
    st.markdown("### 🔥 Postweld Heat Treatment (PWHT)")
    col1, col2, col3 = st.columns(3)
    with col1:
        pwht_temp = st.text_input("Temperature:", placeholder="e.g. N/A or 620°C")
    with col2:
        pwht_time = st.text_input("Time at Temperature:", placeholder="e.g. N/A or 1 hour")
    with col3:
        pwht_other = st.text_input("Other:", placeholder="e.g. Heating rate 150°C/hr")

    st.divider()

    # ── SECTION 5: PROCEDURE (Multi-pass support) ──
    st.markdown("### ⚡ Procedure")

    # Number of passes selector
    col1, col2 = st.columns([1, 5])
    with col1:
        num_passes = st.selectbox("Number of Passes:", list(range(1, 9)), index=0)
    with col2:
        st.write("")  # Spacer for alignment

    col1, col2 = st.columns(2)
    with col1:
        weld_layers = st.text_input("Weld Layer(s):", placeholder="e.g. All / 1 / 2-n")
        weld_passes = st.text_input("Weld Pass(es):", placeholder="e.g. All / Root / Fill")
        process = st.selectbox("Process:", ["GMAW", "FCAW", "SAW", "Other"], index=0)
        process_type = st.selectbox("Type:",
            ["Semiautomatic", "Mechanized", "Automatic", "Manual"])
        position = st.selectbox("Position:",
            ["Flat (1G/1F)", "Horizontal (2G/2F)", "Vertical-Up (3G/3F)",
             "Overhead (4G/4F)", "Fixed Pipe (5G)", "All (6G)"])
        vertical_prog = st.selectbox("Vertical Progression:",
            ["Uphill (Up)", "Downhill (Down)", "N/A"])

    with col2:
        filler_spec = st.text_input("Filler Metal (AWS Spec.):", placeholder="e.g. AWS A5.18")
        aws_class = st.text_input("AWS Classification:", placeholder="e.g. ER70S-3")
        electrode_dia = st.text_input("Diameter:", placeholder="e.g. 1.2 mm / 1.6 mm")
        manufacturer = st.text_input("Manufacturer/Trade Name:", placeholder="e.g. Lincoln Electric")
        preheat_temp = st.text_input("Preheat Temperature:", placeholder="e.g. 10°C [50°F] min")
        interpass_temp = st.text_input("Interpass Temperature:", placeholder="e.g. 250°C [480°F] max")

    st.divider()

    # ── SECTION 6: ELECTRICAL CHARACTERISTICS (Multi-pass columns) ──
    st.markdown("### 🔌 Electrical Characteristics")
    st.markdown(f"**Enter data for {num_passes} weld pass(es):**")

    pass_data_elec = {}
    pass_cols = st.columns(min(num_passes, 4))
    for i in range(num_passes):
        col_idx = i % 4
        with pass_cols[col_idx]:
            st.markdown(f"**Pass {i+1}**")
            pass_data_elec[i+1] = {
                'current_polarity': st.selectbox(f"Current/Polarity (Pass {i+1}):",
                    ["DCEP", "DCEN", "AC", "Pulsed DCEP"], key=f"current_{i}"),
                'transfer_mode': st.selectbox(f"Transfer Mode (Pass {i+1}):",
                    ["Spray", "Pulse", "Short-circuit", "N/A"], key=f"transfer_{i}"),
                'power_source_type': st.selectbox(f"Power Source (Pass {i+1}):",
                    ["CC (Constant Current)", "CV (Constant Voltage)", "Pulsed", "N/A"], key=f"power_{i}"),
                'amps': st.text_input(f"Amps (Pass {i+1}):", placeholder="e.g. 200-250", key=f"amps_{i}"),
                'volts': st.text_input(f"Volts (Pass {i+1}):", placeholder="e.g. 26-30", key=f"volts_{i}"),
                'wire_feed': st.text_input(f"Wire Feed Speed (Pass {i+1}):", placeholder="e.g. 600 mm/min", key=f"wire_{i}"),
                'travel_speed': st.text_input(f"Travel Speed (Pass {i+1}):", placeholder="e.g. 400 mm/min", key=f"travel_{i}"),
                'max_heat_input': st.text_input(f"Max Heat Input (Pass {i+1}):", placeholder="e.g. 2.5 kJ/mm", key=f"heat_{i}"),
            }

    col1, col2 = st.columns(2)
    with col1:
        shielding_gas_comp = st.text_input("Shielding Gas Composition:", placeholder="e.g. Ar+2%O2 or 75%Ar+25%CO2")
    with col2:
        shielding_gas_flow = st.text_input("Shielding Gas Flow Rate:", placeholder="e.g. 20-25 L/min")

    col1, col2 = st.columns(2)
    with col1:
        nozzle_size = st.text_input("Nozzle Size:", placeholder="e.g. 16 mm / 5/8 in")
    with col2:
        contact_tube_dist = st.text_input("Contact Tube to Work Distance:", placeholder="e.g. 19-25 mm")

    st.divider()

    # ── SECTION 7: TECHNIQUE (Multi-pass columns) ──
    st.markdown("### 🛠️ Technique")
    st.markdown(f"**Enter data for {num_passes} weld pass(es):**")

    pass_data_tech = {}
    pass_cols = st.columns(min(num_passes, 4))
    for i in range(num_passes):
        col_idx = i % 4
        with pass_cols[col_idx]:
            st.markdown(f"**Pass {i+1}**")
            pass_data_tech[i+1] = {
                'stringer_weave': st.selectbox(f"Stringer/Weave (Pass {i+1}):",
                    ["Stringer", "Weave", "Both"], key=f"sw_{i}"),
                'multi_single': st.selectbox(f"Multi/Single Pass (Pass {i+1}):",
                    ["Multi-pass", "Single pass"], key=f"ms_{i}"),
                'oscillation': st.selectbox(f"Oscillation (Pass {i+1}):",
                    ["Yes", "No", "N/A"], key=f"osc_{i}"),
                'traverse_length': st.text_input(f"Traverse Length (Pass {i+1}):", placeholder="e.g. 100 mm / N/A", key=f"tl_{i}"),
                'traverse_speed': st.text_input(f"Traverse Speed (Pass {i+1}):", placeholder="e.g. 50 mm/min / N/A", key=f"ts_{i}"),
                'dwell_time': st.text_input(f"Dwell Time (Pass {i+1}):", placeholder="e.g. N/A", key=f"dw_{i}"),
                'num_electrodes': st.text_input(f"Number of Electrodes (Pass {i+1}):", placeholder="e.g. 1 / 2 tandem", key=f"ne_{i}"),
                'peening': st.selectbox(f"Peening (Pass {i+1}):", ["None", "Yes", "N/A"], key=f"peen_{i}"),
                'interpass_cleaning': st.selectbox(f"Interpass Cleaning (Pass {i+1}):",
                    ["Wire brush", "Grinding", "Both", "None"], key=f"ipc_{i}"),
            }

    col1, col2 = st.columns(2)
    with col1:
        other_info = st.text_input("Other Information:", placeholder="e.g. Special requirements, notes")
    with col2:
        st.write("")  # Spacer

    st.divider()

    # ── GENERATE BUTTONS ──
    st.markdown("### 📥 Generate WPS Document")
    col1, col2, col3 = st.columns(3)
    with col1:
        generate_pdf = st.button("📄 Generate PDF", use_container_width=True)
    with col2:
        generate_docx = st.button("📝 Generate Word (.docx)", use_container_width=True)
    with col3:
        generate_latex = st.button("🔬 Generate LaTeX PDF", use_container_width=True, help="Professional typeset PDF via LaTeX")

    # ════════════════════════════════════════════════════════
    # PDF GENERATION (Clean B&W AWS Form J-2 Layout)
    # ════════════════════════════════════════════════════════
    if generate_pdf:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm, inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
            import io

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter,
                rightMargin=0.5*inch, leftMargin=0.5*inch,
                topMargin=0.5*inch, bottomMargin=0.5*inch)

            styles = getSampleStyleSheet()
            story = []

            W = 7.5*inch  # total width for letter paper

            # ── HEADER (ANNEX J / AWS D1.1/D1.1M:2020) ──
            hdr_text_style = ParagraphStyle('HdrText', fontSize=7, fontName='Helvetica',
                textColor=colors.black)
            header_row = Table([
                [Paragraph('ANNEX J', hdr_text_style),
                 Paragraph('AWS D1.1/D1.1M:2020', hdr_text_style)]
            ], colWidths=[W/2, W/2])
            header_row.setStyle(TableStyle([
                ('ALIGN', (0,0), (0,0), 'LEFT'),
                ('ALIGN', (1,0), (1,0), 'RIGHT'),
                ('TOPPADDING', (0,0), (-1,-1), 2),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ]))
            story.append(header_row)
            story.append(Spacer(1, 0.08*inch))

            # ── TITLE ──
            title_style = ParagraphStyle('Title', fontSize=11, fontName='Helvetica-Bold',
                alignment=TA_CENTER)
            story.append(Paragraph('Blank Sample WPS Form (GMAW & FCAW)', title_style))
            story.append(Paragraph('WELDING PROCEDURE SPECIFICATION (WPS)', title_style))
            story.append(Spacer(1, 0.08*inch))

            # ── DOCUMENT HEADER TABLE ──
            def mk_cell(text, bold=False):
                font = 'Helvetica-Bold' if bold else 'Helvetica'
                return Paragraph(str(text) if text else '', ParagraphStyle('C', fontSize=7, fontName=font))

            doc_hdr_rows = [
                [mk_cell('Company Name'), mk_cell(company_name or ''),
                 mk_cell('WPS No.'), mk_cell(wps_number or ''),
                 mk_cell('Rev. No.'), mk_cell(rev_number or ''),
                 mk_cell('Date'), mk_cell(str(wps_date) if wps_date else '')],
                [mk_cell('Authorized by'), mk_cell(authorized_by or ''),
                 mk_cell('Date'), mk_cell(str(auth_date) if auth_date else ''),
                 mk_cell('Supporting PQR(s)'), mk_cell(supporting_pqr or ''),
                 mk_cell('CVN Report'), mk_cell(cvn_report or '')],
            ]
            doc_hdr_tbl = Table(doc_hdr_rows, colWidths=[0.8*inch, 0.9*inch, 0.7*inch, 0.85*inch,
                                                          0.7*inch, 0.75*inch, 0.6*inch, 0.75*inch])
            doc_hdr_tbl.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('TOPPADDING', (0,0), (-1,-1), 2),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
                ('LEFTPADDING', (0,0), (-1,-1), 2),
            ]))
            story.append(doc_hdr_tbl)
            story.append(Spacer(1, 0.05*inch))

            # ── BASE METALS (LEFT) + BASE METAL THICKNESS (RIGHT) ──
            # Left table
            bm_rows = [
                [mk_cell('BASE METALS', True), mk_cell('Specification', True),
                 mk_cell('Type or Grade', True), mk_cell('AWS Group No.', True)],
                [mk_cell('Base Material'), mk_cell(bm_spec or ''),
                 mk_cell(bm_type_grade or ''), mk_cell(bm_aws_group or '')],
                [mk_cell('Welded To'), mk_cell(bm_welded_to or ''),
                 mk_cell(''), mk_cell('')],
                [mk_cell('Backing Material'), mk_cell(bm_backing or ''),
                 mk_cell(''), mk_cell('')],
                [mk_cell('Diameter'), mk_cell(bm_diameter or ''),
                 mk_cell(''), mk_cell('')],
            ]
            bm_tbl = Table(bm_rows, colWidths=[0.9*inch, 1.2*inch, 1.2*inch, 1.2*inch])
            bm_tbl.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('LEFTPADDING', (0,0), (-1,-1), 2),
            ]))

            # Right table
            thick_rows = [
                [mk_cell('BASE METAL THICKNESS', True), mk_cell('As-Welded', True),
                 mk_cell('With PWHT', True)],
                [mk_cell('CJP Groove Welds'), mk_cell(thick_cjp_aw or ''),
                 mk_cell(thick_cjp_pwht or '')],
                [mk_cell('CJP Groove w/CVN'), mk_cell(thick_cjp_cvn_aw or ''),
                 mk_cell(thick_cjp_cvn_pwht or '')],
                [mk_cell('PJP Groove Welds'), mk_cell(thick_pjp_aw or ''),
                 mk_cell(thick_pjp_pwht or '')],
                [mk_cell('Fillet Welds'), mk_cell(thick_fillet_aw or ''),
                 mk_cell(thick_fillet_pwht or '')],
            ]
            thick_tbl = Table(thick_rows, colWidths=[1.5*inch, 1.25*inch, 1.25*inch])
            thick_tbl.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('LEFTPADDING', (0,0), (-1,-1), 2),
            ]))

            # Side-by-side layout
            side_by_side = Table([[bm_tbl, thick_tbl]], colWidths=[3.8*inch, 3.7*inch])
            side_by_side.setStyle(TableStyle([
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(side_by_side)
            story.append(Spacer(1, 0.05*inch))

            # ── JOINT DETAILS (LEFT) + SKETCH (RIGHT) ──
            jd_rows = [
                [mk_cell('JOINT DETAILS', True), mk_cell('', True)],
                [mk_cell('Groove Type'), mk_cell(groove_type or '')],
                [mk_cell('Groove Angle'), mk_cell(groove_angle or '')],
                [mk_cell('Root Opening'), mk_cell(root_opening or '')],
                [mk_cell('Root Face'), mk_cell(root_face or '')],
                [mk_cell('Backgouging'), mk_cell(backgouging or '')],
                [mk_cell('Method'), mk_cell(backgouging_method or '')],
            ]
            jd_tbl = Table(jd_rows, colWidths=[1.5*inch, 2.3*inch])
            jd_tbl.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('LEFTPADDING', (0,0), (-1,-1), 2),
            ]))

            sketch_tbl = Table([
                [mk_cell('JOINT DETAILS (Sketch)', True)],
                [Paragraph('[Sketch area]', ParagraphStyle('S', fontSize=6, fontName='Helvetica-Italic'))],
            ], colWidths=[3.7*inch])
            sketch_tbl.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('MINHEIGHT', (0,1), (-1,1), 1.0*inch),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('LEFTPADDING', (0,0), (-1,-1), 2),
            ]))

            jd_side = Table([[jd_tbl, sketch_tbl]], colWidths=[3.8*inch, 3.7*inch])
            jd_side.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(jd_side)
            story.append(Spacer(1, 0.05*inch))

            # ── POSTWELD HEAT TREATMENT ──
            pwht_rows = [
                [mk_cell('POSTWELD HEAT TREATMENT', True),
                 mk_cell('', True), mk_cell('', True)],
                [mk_cell('Temperature'), mk_cell(pwht_temp or ''), mk_cell('')],
                [mk_cell('Time at Temperature'), mk_cell(pwht_time or ''), mk_cell('')],
                [mk_cell('Other'), mk_cell(pwht_other or ''), mk_cell('')],
            ]
            pwht_tbl = Table(pwht_rows, colWidths=[1.5*inch, 2.5*inch, 3.5*inch])
            pwht_tbl.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('LEFTPADDING', (0,0), (-1,-1), 2),
            ]))
            story.append(pwht_tbl)
            story.append(Spacer(1, 0.05*inch))

            # ── PROCEDURE (Multi-column table) ──
            # Build procedure rows with dynamic columns based on num_passes
            col_widths = [1.2*inch] + [(W - 1.2*inch) / num_passes] * num_passes

            proc_row_labels = [
                'Weld Layer(s)',
                'Weld Pass(es)',
                'Process',
                '  Type (Semiautomatic, Mechanized, etc.)',
                'Position',
                '  Vertical Progression',
                'Filler Metal (AWS Spec.)',
                '  AWS Classification',
                '  Diameter',
                '  Manufacturer/Trade Name',
                'Shielding Gas (Composition)',
                '  Flow Rate',
                '  Nozzle Size',
                'Preheat Temperature',
                '  Interpass Temperature',
                'Electrical Characteristics ' + '—' * (num_passes * 3),
                'Current Type & Polarity',
                '  Transfer Mode',
                '  Power Source Type (cc, cv, etc.)',
                'Amps',
                'Volts',
                'Wire Feed Speed',
                'Travel Speed',
                'Maximum Heat Input',
                'Technique ' + '—' * (num_passes * 3),
                'Stringer or Weave',
                '  Multi or Single Pass (per side)',
                'Oscillation (Mechanized/Automatic)',
                'Traverse Length',
                'Traverse Speed',
                'Dwell Time',
                'Number of Electrodes',
                'Contact Tube to Work Distance',
                'Peening',
                'Interpass Cleaning',
                'Other',
            ]

            proc_rows = []
            # Header row
            header_cells = [mk_cell('PROCEDURE', True)]
            for i in range(1, num_passes + 1):
                header_cells.append(mk_cell(f'Pass {i}', True))
            proc_rows.append(header_cells)

            # Data rows
            for label in proc_row_labels:
                row = [mk_cell(label, 'ELECTRICAL' in label or 'Technique' in label)]
                for p in range(1, num_passes + 1):
                    if label == 'Weld Layer(s)':
                        row.append(mk_cell(weld_layers or ''))
                    elif label == 'Weld Pass(es)':
                        row.append(mk_cell(weld_passes or ''))
                    elif label == 'Process':
                        row.append(mk_cell(process or 'GMAW'))
                    elif label.startswith('  Type'):
                        row.append(mk_cell(process_type or ''))
                    elif label == 'Position':
                        row.append(mk_cell(position or ''))
                    elif label.startswith('  Vertical'):
                        row.append(mk_cell(vertical_prog or ''))
                    elif label == 'Filler Metal (AWS Spec.)':
                        row.append(mk_cell(filler_spec or ''))
                    elif label.startswith('  AWS Classification'):
                        row.append(mk_cell(aws_class or ''))
                    elif label.startswith('  Diameter'):
                        row.append(mk_cell(electrode_dia or ''))
                    elif label.startswith('  Manufacturer'):
                        row.append(mk_cell(manufacturer or ''))
                    elif label == 'Shielding Gas (Composition)':
                        row.append(mk_cell(shielding_gas_comp or ''))
                    elif label.startswith('  Flow Rate'):
                        row.append(mk_cell(shielding_gas_flow or ''))
                    elif label.startswith('  Nozzle'):
                        row.append(mk_cell(nozzle_size or ''))
                    elif label == 'Preheat Temperature':
                        row.append(mk_cell(preheat_temp or ''))
                    elif label.startswith('  Interpass'):
                        row.append(mk_cell(interpass_temp or ''))
                    elif 'Electrical' in label or 'Technique' in label or label.startswith('—'):
                        row.append(mk_cell(''))
                    elif label == 'Current Type & Polarity':
                        row.append(mk_cell(pass_data_elec.get(p, {}).get('current_polarity', '')))
                    elif label.startswith('  Transfer'):
                        row.append(mk_cell(pass_data_elec.get(p, {}).get('transfer_mode', '')))
                    elif label.startswith('  Power'):
                        row.append(mk_cell(pass_data_elec.get(p, {}).get('power_source_type', '')))
                    elif label == 'Amps':
                        row.append(mk_cell(pass_data_elec.get(p, {}).get('amps', '')))
                    elif label == 'Volts':
                        row.append(mk_cell(pass_data_elec.get(p, {}).get('volts', '')))
                    elif label == 'Wire Feed Speed':
                        row.append(mk_cell(pass_data_elec.get(p, {}).get('wire_feed', '')))
                    elif label == 'Travel Speed':
                        row.append(mk_cell(pass_data_elec.get(p, {}).get('travel_speed', '')))
                    elif label == 'Maximum Heat Input':
                        row.append(mk_cell(pass_data_elec.get(p, {}).get('max_heat_input', '')))
                    elif label == 'Stringer or Weave':
                        row.append(mk_cell(pass_data_tech.get(p, {}).get('stringer_weave', '')))
                    elif label.startswith('  Multi'):
                        row.append(mk_cell(pass_data_tech.get(p, {}).get('multi_single', '')))
                    elif label.startswith('Oscillation'):
                        row.append(mk_cell(pass_data_tech.get(p, {}).get('oscillation', '')))
                    elif label == 'Traverse Length':
                        row.append(mk_cell(pass_data_tech.get(p, {}).get('traverse_length', '')))
                    elif label == 'Traverse Speed':
                        row.append(mk_cell(pass_data_tech.get(p, {}).get('traverse_speed', '')))
                    elif label == 'Dwell Time':
                        row.append(mk_cell(pass_data_tech.get(p, {}).get('dwell_time', '')))
                    elif label == 'Number of Electrodes':
                        row.append(mk_cell(pass_data_tech.get(p, {}).get('num_electrodes', '')))
                    elif label.startswith('Contact Tube'):
                        row.append(mk_cell(contact_tube_dist or ''))
                    elif label == 'Peening':
                        row.append(mk_cell(pass_data_tech.get(p, {}).get('peening', '')))
                    elif label.startswith('Interpass Cleaning'):
                        row.append(mk_cell(pass_data_tech.get(p, {}).get('interpass_cleaning', '')))
                    elif label == 'Other':
                        row.append(mk_cell(other_info or ''))
                    else:
                        row.append(mk_cell(''))
                proc_rows.append(row)

            proc_tbl = Table(proc_rows, colWidths=col_widths)
            proc_tbl.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,-1), 6),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('LEFTPADDING', (0,0), (-1,-1), 1),
                ('MINHEIGHT', (0,0), (-1,-1), 0.18*inch),
            ]))
            story.append(proc_tbl)
            story.append(Spacer(1, 0.05*inch))

            # ── FOOTER ──
            footer_style = ParagraphStyle('Footer', fontSize=6, fontName='Helvetica',
                alignment=TA_CENTER)
            footer_text = Paragraph(
                'Form J-2 &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; '
                '(See http://go.aws.org/D1forms)',
                footer_style
            )
            story.append(footer_text)
            story.append(Spacer(1, 0.02*inch))

            footer_bottom = Paragraph(
                'ANNEX J &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; AWS D1.1/D1.1M:2020',
                footer_style
            )
            story.append(footer_bottom)

            doc.build(story)
            buffer.seek(0)

            st.success("✅ PDF generated successfully!")
            st.download_button(
                label="📥 Download WPS PDF",
                data=buffer,
                file_name=f"WPS_{wps_number}_{wps_date}.pdf",
                mime="application/pdf",
                use_container_width=True
            )

        except Exception as e:
            st.error(f"❌ Error generating PDF: {str(e)}")

    # ════════════════════════════════════════════════════════
    # DOCX GENERATION (Clean B&W AWS Form J-2 Layout)
    # ════════════════════════════════════════════════════════
    if generate_docx:
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import io

            docx_buffer = io.BytesIO()
            document = DocxDocument()

            section = document.sections[0]
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)

            def set_cell_borders(cell, **kwargs):
                tcPr = cell._tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    cell._element.append(tcPr)
                tcBorders = OxmlElement('w:tcBorders')
                for edge in ('top', 'left', 'bottom', 'right'):
                    edge_data = kwargs.get(edge)
                    if edge_data is not None:
                        edge_el = OxmlElement(f'w:{edge}')
                        edge_el.set(qn('w:val'), 'single')
                        edge_el.set(qn('w:sz'), '4')
                        edge_el.set(qn('w:space'), '0')
                        edge_el.set(qn('w:color'), '000000')
                        tcBorders.append(edge_el)
                tcPr.append(tcBorders)

            def set_cell_bg(cell, hex_color):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), hex_color)
                tcPr.append(shd)

            # ── Header ──
            hdr_para = document.add_paragraph()
            hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hdr_para.add_run('ANNEX J — AWS D1.1/D1.1M:2020')
            run.font.size = Pt(8)
            run.font.bold = True

            # ── Title ──
            title_para = document.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run('Blank Sample WPS Form (GMAW & FCAW)\n')
            run.font.size = Pt(10)
            run.font.bold = True
            run = title_para.add_run('WELDING PROCEDURE SPECIFICATION (WPS)')
            run.font.size = Pt(10)
            run.font.bold = True

            document.add_paragraph()

            # ── Document Header Table ──
            hdr_tbl = document.add_table(rows=2, cols=8)
            hdr_tbl.style = 'Table Grid'
            for i, label in enumerate(['Company Name', 'WPS No.', 'Rev. No.', 'Date',
                                       'Authorized by', 'Auth. Date', 'Supporting PQR(s)', 'CVN Report']):
                if i < 4:
                    cell = hdr_tbl.rows[0].cells[i]
                    cell.text = label
                else:
                    cell = hdr_tbl.rows[1].cells[i - 4]
                    cell.text = label
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(7)

            hdr_vals = [company_name, wps_number, rev_number, str(wps_date),
                       authorized_by, str(auth_date), supporting_pqr, cvn_report]
            for i, val in enumerate(hdr_vals):
                if i < 4:
                    cell = hdr_tbl.rows[0].cells[i]
                    cell.text = str(val) if val else ''
                else:
                    cell = hdr_tbl.rows[1].cells[i - 4]
                    cell.text = str(val) if val else ''
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(7)

            document.add_paragraph()

            # ── Base Metals ──
            document.add_paragraph('BASE METALS', style='Heading 2')
            bm_tbl = document.add_table(rows=5, cols=4)
            bm_tbl.style = 'Table Grid'
            headers = ['', 'Specification', 'Type or Grade', 'AWS Group No.']
            for i, h in enumerate(headers):
                cell = bm_tbl.rows[0].cells[i]
                cell.text = h
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(7)

            data = [
                ['Base Material', bm_spec, bm_type_grade, bm_aws_group],
                ['Welded To', bm_welded_to, '', ''],
                ['Backing Material', bm_backing, '', ''],
                ['Diameter', bm_diameter, '', ''],
            ]
            for i, row_data in enumerate(data):
                for j, val in enumerate(row_data):
                    cell = bm_tbl.rows[i + 1].cells[j]
                    cell.text = str(val) if val else ''
                    if j == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(7)

            document.add_paragraph()

            # ── Base Metal Thickness ──
            thick_tbl = document.add_table(rows=5, cols=3)
            thick_tbl.style = 'Table Grid'
            headers = ['Base Metal Thickness', 'As-Welded', 'With PWHT']
            for i, h in enumerate(headers):
                cell = thick_tbl.rows[0].cells[i]
                cell.text = h
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(7)

            thick_data = [
                ['CJP Groove Welds', thick_cjp_aw, thick_cjp_pwht],
                ['CJP Groove w/CVN', thick_cjp_cvn_aw, thick_cjp_cvn_pwht],
                ['PJP Groove Welds', thick_pjp_aw, thick_pjp_pwht],
                ['Fillet Welds', thick_fillet_aw, thick_fillet_pwht],
            ]
            for i, row_data in enumerate(thick_data):
                for j, val in enumerate(row_data):
                    cell = thick_tbl.rows[i + 1].cells[j]
                    cell.text = str(val) if val else ''
                    if j == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(7)

            document.add_paragraph()

            # ── Joint Details ──
            document.add_paragraph('JOINT DETAILS', style='Heading 2')
            jd_tbl = document.add_table(rows=7, cols=2)
            jd_tbl.style = 'Table Grid'
            jd_data = [
                ['Groove Type', groove_type],
                ['Groove Angle', groove_angle],
                ['Root Opening', root_opening],
                ['Root Face', root_face],
                ['Backgouging', backgouging],
                ['Method', backgouging_method],
            ]
            for i, row_data in enumerate(jd_data):
                cell = jd_tbl.rows[i].cells[0]
                cell.text = row_data[0]
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(7)
                cell = jd_tbl.rows[i].cells[1]
                cell.text = str(row_data[1]) if row_data[1] else ''
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(7)

            document.add_paragraph()

            # ── PWHT ──
            document.add_paragraph('POSTWELD HEAT TREATMENT', style='Heading 2')
            pwht_tbl = document.add_table(rows=4, cols=2)
            pwht_tbl.style = 'Table Grid'
            pwht_data = [
                ['Temperature', pwht_temp],
                ['Time at Temperature', pwht_time],
                ['Other', pwht_other],
            ]
            for i, row_data in enumerate(pwht_data):
                cell = pwht_tbl.rows[i].cells[0]
                cell.text = row_data[0]
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(7)
                cell = pwht_tbl.rows[i].cells[1]
                cell.text = str(row_data[1]) if row_data[1] else ''
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(7)

            document.add_paragraph()

            # ── Procedure ──
            document.add_paragraph('PROCEDURE', style='Heading 2')
            proc_tbl = document.add_table(rows=1, cols=2)
            proc_tbl.style = 'Table Grid'
            proc_data = [
                ['Weld Layer(s)', weld_layers],
                ['Weld Pass(es)', weld_passes],
                ['Process', process],
                ['Type', process_type],
                ['Position', position],
                ['Vertical Progression', vertical_prog],
                ['Filler Metal (AWS Spec.)', filler_spec],
                ['AWS Classification', aws_class],
                ['Diameter', electrode_dia],
                ['Manufacturer/Trade Name', manufacturer],
                ['Shielding Gas Composition', shielding_gas_comp],
                ['Flow Rate', shielding_gas_flow],
                ['Nozzle Size', nozzle_size],
                ['Preheat Temperature', preheat_temp],
                ['Interpass Temperature', interpass_temp],
                ['Contact Tube to Work Distance', contact_tube_dist],
            ]
            for _ in proc_data:
                proc_tbl.add_row()
            for i, row_data in enumerate(proc_data):
                cell = proc_tbl.rows[i].cells[0]
                cell.text = row_data[0]
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(7)
                cell = proc_tbl.rows[i].cells[1]
                cell.text = str(row_data[1]) if row_data[1] else ''
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(7)

            document.add_paragraph()

            # ── Electrical Characteristics (per pass) ──
            document.add_paragraph('ELECTRICAL CHARACTERISTICS', style='Heading 2')
            for p in range(1, num_passes + 1):
                document.add_paragraph(f'Pass {p}:', style='List Bullet')
                elec_data = pass_data_elec.get(p, {})
                p_data = [
                    ['Current Type & Polarity', elec_data.get('current_polarity', '')],
                    ['Transfer Mode', elec_data.get('transfer_mode', '')],
                    ['Power Source Type', elec_data.get('power_source_type', '')],
                    ['Amps', elec_data.get('amps', '')],
                    ['Volts', elec_data.get('volts', '')],
                    ['Wire Feed Speed', elec_data.get('wire_feed', '')],
                    ['Travel Speed', elec_data.get('travel_speed', '')],
                    ['Maximum Heat Input', elec_data.get('max_heat_input', '')],
                ]
                elec_tbl = document.add_table(rows=len(p_data), cols=2)
                elec_tbl.style = 'Table Grid'
                for i, row_data in enumerate(p_data):
                    cell = elec_tbl.rows[i].cells[0]
                    cell.text = row_data[0]
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                            run.font.size = Pt(7)
                    cell = elec_tbl.rows[i].cells[1]
                    cell.text = str(row_data[1]) if row_data[1] else ''
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(7)

            document.add_paragraph()

            # ── Technique (per pass) ──
            document.add_paragraph('TECHNIQUE', style='Heading 2')
            for p in range(1, num_passes + 1):
                document.add_paragraph(f'Pass {p}:', style='List Bullet')
                tech_data = pass_data_tech.get(p, {})
                t_data = [
                    ['Stringer or Weave', tech_data.get('stringer_weave', '')],
                    ['Multi or Single Pass', tech_data.get('multi_single', '')],
                    ['Oscillation', tech_data.get('oscillation', '')],
                    ['Traverse Length', tech_data.get('traverse_length', '')],
                    ['Traverse Speed', tech_data.get('traverse_speed', '')],
                    ['Dwell Time', tech_data.get('dwell_time', '')],
                    ['Number of Electrodes', tech_data.get('num_electrodes', '')],
                    ['Peening', tech_data.get('peening', '')],
                    ['Interpass Cleaning', tech_data.get('interpass_cleaning', '')],
                ]
                tech_tbl = document.add_table(rows=len(t_data), cols=2)
                tech_tbl.style = 'Table Grid'
                for i, row_data in enumerate(t_data):
                    cell = tech_tbl.rows[i].cells[0]
                    cell.text = row_data[0]
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                            run.font.size = Pt(7)
                    cell = tech_tbl.rows[i].cells[1]
                    cell.text = str(row_data[1]) if row_data[1] else ''
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(7)

            document.add_paragraph()

            # ── Footer ──
            footer_para = document.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer_para.add_run(
                'Form J-2 | AWS D1.1/D1.1M:2020 Annex J | '
                'Generated by Welding Engineering Toolbox')
            run.font.size = Pt(6)
            run.font.italic = True

            document.save(docx_buffer)
            docx_buffer.seek(0)

            st.success("✅ Word document generated successfully!")
            st.download_button(
                label="📥 Download WPS Word (.docx)",
                data=docx_buffer,
                file_name=f"WPS_{wps_number}_{wps_date}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        except Exception as e:
            st.error(f"❌ Error generating Word document: {str(e)}")

    # ════════════════════════════════════════════════════════
    # LATEX PDF GENERATION (Clean B&W AWS Form J-2 Layout)
    # ════════════════════════════════════════════════════════
    if generate_latex:
        try:
            import requests, io

            def _esc(text):
                """Escape special LaTeX characters."""
                s = str(text) if text else "---"
                for ch in ['\\', '&', '%', '$', '#', '_', '{', '}', '~', '^']:
                    s = s.replace(ch, '\\' + ch)
                return s

            # ── Build LaTeX source (LETTER, B&W, clean form layout) ──
            latex_src = r"""\documentclass[letterpaper,10pt]{article}
\usepackage[margin=0.5in]{geometry}
\usepackage{booktabs,tabularx,colortbl,xcolor,array,graphicx}
\usepackage[T1]{fontenc}
\usepackage{lmodern}
\usepackage{fancyhdr}

% ── Page style ──
\pagestyle{fancy}
\fancyhf{}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}
\fancyfoot[C]{\scriptsize Form J-2 \quad AWS D1.1/D1.1M:2020 --- Annex J}

\setlength{\parindent}{0pt}
\setlength{\tabcolsep}{3pt}
\renewcommand{\arraystretch}{1.2}

\begin{document}

% ══════════════════════════  HEADER  ══════════════════════════
\noindent
\begin{tabularx}{\textwidth}{|X|X|}
\hline
\textbf{ANNEX J} & \textbf{AWS D1.1/D1.1M:2020} \\
\hline
\end{tabularx}

\vspace{6pt}

% ══════════════════════════  TITLE  ══════════════════════════
\begin{center}
\textbf{Blank Sample WPS Form (GMAW \& FCAW)}\\
\textbf{WELDING PROCEDURE SPECIFICATION (WPS)}
\end{center}

\vspace{6pt}

% ══════════════════════════  DOCUMENT HEADER  ══════════════════════════
\noindent
\begin{tabularx}{\textwidth}{|X|X|X|X|X|X|X|X|}
\hline
""" + f"""\\textbf{{Company Name}} & {_esc(company_name)} & \\textbf{{WPS No.}} & {_esc(wps_number)} & \\textbf{{Rev. No.}} & {_esc(rev_number)} & \\textbf{{Date}} & {_esc(str(wps_date))} \\\\
\\hline
\\textbf{{Authorized by}} & {_esc(authorized_by)} & \\textbf{{Date}} & {_esc(str(auth_date))} & \\textbf{{Supporting PQR(s)}} & {_esc(supporting_pqr)} & \\textbf{{CVN Report}} & {_esc(cvn_report)} \\\\
\\hline
\\end{{tabularx}}

\vspace{{4pt}}

% ══════════════════════════  BASE METALS  ══════════════════════════
\\noindent
\\begin{{tabularx}}{{\\textwidth}}{{|X|X|X|X|}}
\\hline
\\multicolumn{{4}}{{|X|}}{{\\textbf{{BASE METALS}}}} \\\\
\\hline
\\textbf{{Specification}} & {_esc(bm_spec)} & \\textbf{{Type or Grade}} & {_esc(bm_type_grade)} \\\\
\\hline
\\textbf{{AWS Group No.}} & {_esc(bm_aws_group)} & \\textbf{{Welded To}} & {_esc(bm_welded_to)} \\\\
\\hline
\\textbf{{Backing Material}} & {_esc(bm_backing)} & \\textbf{{Diameter}} & {_esc(bm_diameter)} \\\\
\\hline
\\end{{tabularx}}

\\noindent
\\begin{{tabularx}}{{\\textwidth}}{{|X|X|X|}}
\\hline
\\multicolumn{{3}}{{|X|}}{{\\textbf{{BASE METAL THICKNESS}}}} \\\\
\\hline
\\textbf{{Weld Type}} & \\textbf{{As-Welded}} & \\textbf{{With PWHT}} \\\\
\\hline
CJP Groove Welds & {_esc(thick_cjp_aw)} & {_esc(thick_cjp_pwht)} \\\\
\\hline
CJP Groove w/CVN & {_esc(thick_cjp_cvn_aw)} & {_esc(thick_cjp_cvn_pwht)} \\\\
\\hline
PJP Groove Welds & {_esc(thick_pjp_aw)} & {_esc(thick_pjp_pwht)} \\\\
\\hline
Fillet Welds & {_esc(thick_fillet_aw)} & {_esc(thick_fillet_pwht)} \\\\
\\hline
\\end{{tabularx}}

\\vspace{{4pt}}

% ══════════════════════════  JOINT DETAILS  ══════════════════════════
\\noindent
\\begin{{tabularx}}{{0.48\\textwidth}}{{|X|X|}}
\\hline
\\multicolumn{{2}}{{|X|}}{{\\textbf{{JOINT DETAILS}}}} \\\\
\\hline
\\textbf{{Groove Type}} & {_esc(groove_type)} \\\\
\\hline
\\textbf{{Groove Angle}} & {_esc(groove_angle)} \\\\
\\hline
\\textbf{{Root Opening}} & {_esc(root_opening)} \\\\
\\hline
\\textbf{{Root Face}} & {_esc(root_face)} \\\\
\\hline
\\textbf{{Backgouging}} & {_esc(backgouging)} \\\\
\\hline
\\textbf{{Method}} & {_esc(backgouging_method)} \\\\
\\hline
\\end{{tabularx}}
\\hfill
\\begin{{tabularx}}{{0.48\\textwidth}}{{|X|}}
\\hline
\\textbf{{JOINT DETAILS (Sketch)}} \\\\
\\hline
\\mbox{{}} \\\\
\\mbox{{}} \\\\
\\mbox{{}} \\\\
\\mbox{{}} \\\\
\\mbox{{}} \\\\
\\mbox{{}} \\\\
\\hline
\\end{{tabularx}}

\\vspace{{4pt}}

% ══════════════════════════  POSTWELD HEAT TREATMENT  ══════════════════════════
\\noindent
\\begin{{tabularx}}{{\\textwidth}}{{|X|X|X|}}
\\hline
\\multicolumn{{3}}{{|X|}}{{\\textbf{{POSTWELD HEAT TREATMENT}}}} \\\\
\\hline
\\textbf{{Temperature}} & \\textbf{{Time at Temperature}} & \\textbf{{Other}} \\\\
\\hline
{_esc(pwht_temp)} & {_esc(pwht_time)} & {_esc(pwht_other)} \\\\
\\hline
\\end{{tabularx}}

\\vspace{{4pt}}

% ══════════════════════════  PROCEDURE (Multi-pass)  ══════════════════════════
\\noindent
\\textbf{{PROCEDURE}} (Weld Layer(s): {_esc(weld_layers)}, Pass(es): {_esc(weld_passes)})

\\smallskip

\\noindent
\\begin{{tabularx}}{{\\textwidth}}{{|X|X|X|X|}}
\\hline
\\multicolumn{{4}}{{|X|}}{{\\textbf{{PROCEDURE DETAILS}}}} \\\\
\\hline
\\textbf{{Parameter}} & \\textbf{{Value}} & \\textbf{{Parameter}} & \\textbf{{Value}} \\\\
\\hline
Process & {_esc(process)} & Vertical Progression & {_esc(vertical_prog)} \\\\
\\hline
Type & {_esc(process_type)} & Position & {_esc(position)} \\\\
\\hline
Filler Metal (AWS) & {_esc(filler_spec)} & AWS Classification & {_esc(aws_class)} \\\\
\\hline
Diameter & {_esc(electrode_dia)} & Manufacturer & {_esc(manufacturer)} \\\\
\\hline
Shielding Gas & {_esc(shielding_gas_comp)} & Flow Rate & {_esc(shielding_gas_flow)} \\\\
\\hline
Nozzle Size & {_esc(nozzle_size)} & Preheat Temp. & {_esc(preheat_temp)} \\\\
\\hline
Interpass Temp. & {_esc(interpass_temp)} & Contact Tube to Work & {_esc(contact_tube_dist)} \\\\
\\hline
\\end{{tabularx}}

\\vspace{{4pt}}

% ══════════════════════════  ELECTRICAL (Per-pass)  ══════════════════════════
\\noindent
\\textbf{{ELECTRICAL CHARACTERISTICS}}

\\smallskip
\\noindent
\\begin{{tabularx}}{{\\textwidth}}{{|X|X|X|X|X|X|}}
\\hline
\\textbf{{Parameter}} & \\textbf{{Pass 1}} & \\textbf{{Pass 2}} & \\textbf{{Pass 3}} & \\textbf{{Pass 4}} & \\textbf{{Pass 5}} \\\\
\\hline
Current Type & {_esc(pass_data_elec.get(1, {}).get('current_polarity', ''))} & {_esc(pass_data_elec.get(2, {}).get('current_polarity', ''))} & {_esc(pass_data_elec.get(3, {}).get('current_polarity', ''))} & {_esc(pass_data_elec.get(4, {}).get('current_polarity', ''))} & {_esc(pass_data_elec.get(5, {}).get('current_polarity', ''))} \\\\
\\hline
Transfer Mode & {_esc(pass_data_elec.get(1, {}).get('transfer_mode', ''))} & {_esc(pass_data_elec.get(2, {}).get('transfer_mode', ''))} & {_esc(pass_data_elec.get(3, {}).get('transfer_mode', ''))} & {_esc(pass_data_elec.get(4, {}).get('transfer_mode', ''))} & {_esc(pass_data_elec.get(5, {}).get('transfer_mode', ''))} \\\\
\\hline
Amps & {_esc(pass_data_elec.get(1, {}).get('amps', ''))} & {_esc(pass_data_elec.get(2, {}).get('amps', ''))} & {_esc(pass_data_elec.get(3, {}).get('amps', ''))} & {_esc(pass_data_elec.get(4, {}).get('amps', ''))} & {_esc(pass_data_elec.get(5, {}).get('amps', ''))} \\\\
\\hline
Volts & {_esc(pass_data_elec.get(1, {}).get('volts', ''))} & {_esc(pass_data_elec.get(2, {}).get('volts', ''))} & {_esc(pass_data_elec.get(3, {}).get('volts', ''))} & {_esc(pass_data_elec.get(4, {}).get('volts', ''))} & {_esc(pass_data_elec.get(5, {}).get('volts', ''))} \\\\
\\hline
Wire Feed Speed & {_esc(pass_data_elec.get(1, {}).get('wire_feed', ''))} & {_esc(pass_data_elec.get(2, {}).get('wire_feed', ''))} & {_esc(pass_data_elec.get(3, {}).get('wire_feed', ''))} & {_esc(pass_data_elec.get(4, {}).get('wire_feed', ''))} & {_esc(pass_data_elec.get(5, {}).get('wire_feed', ''))} \\\\
\\hline
Travel Speed & {_esc(pass_data_elec.get(1, {}).get('travel_speed', ''))} & {_esc(pass_data_elec.get(2, {}).get('travel_speed', ''))} & {_esc(pass_data_elec.get(3, {}).get('travel_speed', ''))} & {_esc(pass_data_elec.get(4, {}).get('travel_speed', ''))} & {_esc(pass_data_elec.get(5, {}).get('travel_speed', ''))} \\\\
\\hline
Max Heat Input & {_esc(pass_data_elec.get(1, {}).get('max_heat_input', ''))} & {_esc(pass_data_elec.get(2, {}).get('max_heat_input', ''))} & {_esc(pass_data_elec.get(3, {}).get('max_heat_input', ''))} & {_esc(pass_data_elec.get(4, {}).get('max_heat_input', ''))} & {_esc(pass_data_elec.get(5, {}).get('max_heat_input', ''))} \\\\
\\hline
\\end{{tabularx}}

\\vspace{{4pt}}

% ══════════════════════════  TECHNIQUE (Per-pass)  ══════════════════════════
\\noindent
\\textbf{{TECHNIQUE}}

\\smallskip
\\noindent
\\begin{{tabularx}}{{\\textwidth}}{{|X|X|X|X|X|X|}}
\\hline
\\textbf{{Parameter}} & \\textbf{{Pass 1}} & \\textbf{{Pass 2}} & \\textbf{{Pass 3}} & \\textbf{{Pass 4}} & \\textbf{{Pass 5}} \\\\
\\hline
Stringer/Weave & {_esc(pass_data_tech.get(1, {}).get('stringer_weave', ''))} & {_esc(pass_data_tech.get(2, {}).get('stringer_weave', ''))} & {_esc(pass_data_tech.get(3, {}).get('stringer_weave', ''))} & {_esc(pass_data_tech.get(4, {}).get('stringer_weave', ''))} & {_esc(pass_data_tech.get(5, {}).get('stringer_weave', ''))} \\\\
\\hline
Multi/Single Pass & {_esc(pass_data_tech.get(1, {}).get('multi_single', ''))} & {_esc(pass_data_tech.get(2, {}).get('multi_single', ''))} & {_esc(pass_data_tech.get(3, {}).get('multi_single', ''))} & {_esc(pass_data_tech.get(4, {}).get('multi_single', ''))} & {_esc(pass_data_tech.get(5, {}).get('multi_single', ''))} \\\\
\\hline
Oscillation & {_esc(pass_data_tech.get(1, {}).get('oscillation', ''))} & {_esc(pass_data_tech.get(2, {}).get('oscillation', ''))} & {_esc(pass_data_tech.get(3, {}).get('oscillation', ''))} & {_esc(pass_data_tech.get(4, {}).get('oscillation', ''))} & {_esc(pass_data_tech.get(5, {}).get('oscillation', ''))} \\\\
\\hline
Traverse Length & {_esc(pass_data_tech.get(1, {}).get('traverse_length', ''))} & {_esc(pass_data_tech.get(2, {}).get('traverse_length', ''))} & {_esc(pass_data_tech.get(3, {}).get('traverse_length', ''))} & {_esc(pass_data_tech.get(4, {}).get('traverse_length', ''))} & {_esc(pass_data_tech.get(5, {}).get('traverse_length', ''))} \\\\
\\hline
Peening & {_esc(pass_data_tech.get(1, {}).get('peening', ''))} & {_esc(pass_data_tech.get(2, {}).get('peening', ''))} & {_esc(pass_data_tech.get(3, {}).get('peening', ''))} & {_esc(pass_data_tech.get(4, {}).get('peening', ''))} & {_esc(pass_data_tech.get(5, {}).get('peening', ''))} \\\\
\\hline
Interpass Cleaning & {_esc(pass_data_tech.get(1, {}).get('interpass_cleaning', ''))} & {_esc(pass_data_tech.get(2, {}).get('interpass_cleaning', ''))} & {_esc(pass_data_tech.get(3, {}).get('interpass_cleaning', ''))} & {_esc(pass_data_tech.get(4, {}).get('interpass_cleaning', ''))} & {_esc(pass_data_tech.get(5, {}).get('interpass_cleaning', ''))} \\\\
\\hline
\\end{{tabularx}}

\\vspace{{10pt}}

\\end{{document}}
"""

            with st.spinner("⏳ Compiling LaTeX via latex.ytotech.com …"):
                resp = requests.post(
                    "https://latex.ytotech.com/builds/sync",
                    json={
                        "compiler": "pdflatex",
                        "resources": [
                            {
                                "main": True,
                                "content": latex_src
                            }
                        ]
                    },
                    timeout=60
                )

            if resp.status_code in (200, 201) and resp.content[:5] == b"%PDF-":
                st.success("✅ LaTeX PDF generated successfully!")
                st.download_button(
                    label="📥 Download LaTeX WPS PDF",
                    data=resp.content,
                    file_name=f"WPS_{wps_number}_{wps_date}_latex.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            else:
                error_detail = resp.text[:1000] if resp.text else "No response body"
                st.error(f"❌ LaTeX compilation failed (HTTP {resp.status_code}).\n\n{error_detail}")
                with st.expander("📄 View LaTeX source for debugging"):
                    st.code(latex_src, language="latex")

        except requests.exceptions.Timeout:
            st.error("❌ LaTeX compilation timed out. The service may be busy — please try again.")
        except Exception as e:
            st.error(f"❌ Error generating LaTeX PDF: {str(e)}")
            import traceback
            with st.expander("Debug info"):
                st.code(traceback.format_exc())

st.divider()
st.caption("📖 Reference: AWS D1.1/D1.1M:2025 — Clause 5 | For educational reference only. Always verify against the official published standard.")